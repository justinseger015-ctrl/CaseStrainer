# DEPRECATED: This file is no longer used in the main pipeline. See src/document_processing.py and src/api/services/citation_service.py for the current implementation.

"""
Unified document processing module for CaseStrainer.
Handles text extraction from various file formats and input types.

LEGACY VERSION - This is a backup of the previous document_processing.py.
The enhanced version is now in document_processing.py.
"""

import os
import logging
import tempfile
from pathlib import Path
from typing import Dict, Any, Optional, Tuple, List
import requests
from urllib.parse import urlparse
import re
import time

# Import text extraction libraries
try:
    from bs4 import BeautifulSoup
    BEAUTIFULSOUP_AVAILABLE = True
except ImportError:
    BEAUTIFULSOUP_AVAILABLE = False
    logging.warning("BeautifulSoup not available - HTML/HTM processing will be limited")

try:
    from striprtf.striprtf import rtf_to_text
    STRIPRTF_AVAILABLE = True
except ImportError:
    STRIPRTF_AVAILABLE = False
    logging.warning("striprtf not available - RTF processing will not work")

# Import existing processing modules
from src.pdf_handler import PDFHandler, PDFExtractionConfig, PDFExtractionMethod
# DEPRECATED: from src.citation_extractor import CitationExtractor
# DEPRECATED: from src.complex_citation_integration import ComplexCitationIntegrator, format_complex_citation_for_frontend

logger = logging.getLogger(__name__)


def extract_text_from_html_file(file_path: str) -> str:
    """
    Extract text from HTML/HTM files using BeautifulSoup.
    
    Args:
        file_path: Path to the HTML/HTM file
        
    Returns:
        Extracted text content
    """
    if not BEAUTIFULSOUP_AVAILABLE:
        raise ImportError("BeautifulSoup is required for HTML/HTM processing")
    
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        # Parse HTML with BeautifulSoup
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Extract text with proper spacing
        text = soup.get_text(separator='\n', strip=True)
        
        # Clean up extra whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        logger.info(f"Successfully extracted {len(text)} characters from HTML file")
        return text
        
    except UnicodeDecodeError:
        # Try with different encoding
        with open(file_path, 'r', encoding='latin-1') as f:
            html_content = f.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        for script in soup(["script", "style"]):
            script.decompose()
        
        text = soup.get_text(separator='\n', strip=True)
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        logger.info(f"Successfully extracted {len(text)} characters from HTML file using latin-1 encoding")
        return text
        
    except Exception as e:
        logger.error(f"Error extracting text from HTML file: {str(e)}")
        raise


def extract_text_from_rtf_file(file_path: str) -> str:
    """
    Extract text from RTF files using striprtf.
    
    Args:
        file_path: Path to the RTF file
        
    Returns:
        Extracted text content
    """
    if not STRIPRTF_AVAILABLE:
        raise ImportError("striprtf is required for RTF processing")
    
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            rtf_content = f.read()
        
        # Convert RTF to plain text
        text = rtf_to_text(rtf_content)
        
        # Clean up the text
        text = text.strip()
        
        logger.info(f"Successfully extracted {len(text)} characters from RTF file")
        return text
        
    except UnicodeDecodeError:
        # Try with different encoding
        with open(file_path, 'r', encoding='latin-1') as f:
            rtf_content = f.read()
        
        text = rtf_to_text(rtf_content)
        text = text.strip()
        
        logger.info(f"Successfully extracted {len(text)} characters from RTF file using latin-1 encoding")
        return text
        
    except Exception as e:
        logger.error(f"Error extracting text from RTF file: {str(e)}")
        raise


def extract_text_from_odt_file(file_path: str) -> str:
    """
    Extract text from ODT files using python-docx or alternative method.
    
    Args:
        file_path: Path to the ODT file
        
    Returns:
        Extracted text content
    """
    try:
        # Try using python-docx first (works for some ODT files)
        import docx
        doc = docx.Document(file_path)
        paragraphs = [paragraph.text for paragraph in doc.paragraphs]
        text = "\n".join(paragraphs)
        
        logger.info(f"Successfully extracted {len(text)} characters from ODT file using python-docx")
        return text
        
    except Exception as e:
        logger.warning(f"python-docx failed for ODT file: {str(e)}")
        
        # Fallback: try to read as text file
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            logger.info(f"Successfully extracted {len(text)} characters from ODT file as text")
            return text
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as f:
                text = f.read()
            logger.info(f"Successfully extracted {len(text)} characters from ODT file as text using latin-1")
            return text
        except Exception as e2:
            logger.error(f"Error extracting text from ODT file: {str(e2)}")
            raise


def extract_text_from_file(file_path: str) -> str:
    """
    Extract text from a file based on its extension.
    
    Args:
        file_path: Path to the file
        
    Returns:
        Extracted text content
    """
    file_extension = Path(file_path).suffix.lower()
    
    try:
        if file_extension == '.pdf':
            logger.info("Processing PDF file")
            handler = PDFHandler(PDFExtractionConfig(
                preferred_method=PDFExtractionMethod.PDFMINER,
                use_fallback=False,
                timeout=30,
                clean_text=True
            ))
            text = handler.extract_text(file_path)
            if text.startswith("Error:"):
                raise Exception(text)
            return text
            
        elif file_extension in ['.doc', '.docx']:
            logger.info("Processing Word document")
            import docx
            doc = docx.Document(file_path)
            paragraphs = [paragraph.text for paragraph in doc.paragraphs]
            text = "\n".join(paragraphs)
            return text
            
        elif file_extension in ['.html', '.htm']:
            logger.info("Processing HTML/HTM file")
            return extract_text_from_html_file(file_path)
            
        elif file_extension == '.rtf':
            logger.info("Processing RTF file")
            return extract_text_from_rtf_file(file_path)
            
        elif file_extension == '.odt':
            logger.info("Processing ODT file")
            return extract_text_from_odt_file(file_path)
            
        elif file_extension == '.txt':
            logger.info("Processing text file")
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except UnicodeDecodeError:
                with open(file_path, 'r', encoding='latin-1') as f:
                    return f.read()
                    
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
            
    except Exception as e:
        logger.error(f"Error extracting text from file {file_path}: {str(e)}")
        raise


def extract_text_from_url(url: str) -> str:
    """
    Extract text from a URL, handling different content types including PDFs and other document formats.
    
    Args:
        url: URL to extract text from
        
    Returns:
        Extracted text content
    """
    try:
        logger.info(f"Fetching content from URL: {url}")
        
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        
        response = requests.get(url, headers=headers, timeout=30, stream=True)
        response.raise_for_status()
        
        content_type = response.headers.get('content-type', '').lower()
        logger.info(f"Content type of URL {url}: {content_type}")
        
        # Mapping of content types to file extensions for temporary files
        content_type_to_extension = {
            'application/pdf': '.pdf',
            'application/msword': '.doc',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
            'application/rtf': '.rtf',
            'application/vnd.oasis.opendocument.text': '.odt'
        }
        
        # Check if the content type corresponds to a supported document format
        for ct, ext in content_type_to_extension.items():
            if ct in content_type:
                # Handle document content by downloading to a temporary file and processing as a file
                import tempfile
                with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as temp_file:
                    for chunk in response.iter_content(chunk_size=8192):
                        temp_file.write(chunk)
                    temp_file_path = temp_file.name
                logger.info(f"Downloaded content to temporary file: {temp_file_path}")
                
                try:
                    text = extract_text_from_file(temp_file_path)
                    logger.info(f"Successfully extracted {len(text)} characters from {ct} URL")
                    logger.info(f"Sample of extracted text (first 500 characters): {text[:500]}...")
                    return text
                finally:
                    os.unlink(temp_file_path)  # Clean up temporary file
                    logger.info(f"Deleted temporary file: {temp_file_path}")
        
        if 'text/html' in content_type:
            # Parse HTML content
            if BEAUTIFULSOUP_AVAILABLE:
                soup = BeautifulSoup(response.text, 'html.parser')
                for script in soup(["script", "style"]):
                    script.decompose()
                text = soup.get_text(separator='\n', strip=True)
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = '\n'.join(chunk for chunk in chunks if chunk)
            else:
                # Fallback to raw text
                text = response.text
        else:
            # Treat as plain text
            text = response.text
        
        logger.info(f"Successfully extracted {len(text)} characters from URL")
        logger.info(f"Sample of extracted text (first 500 characters): {text[:500]}...")
        return text
        
    except Exception as e:
        logger.error(f"Error extracting text from URL {url}: {str(e)}")
        raise


def process_document(
    content: str = None,
    file_path: str = None,
    url: str = None,
    extract_case_names: bool = True
) -> Dict[str, Any]:
    """
    Process a document and extract citations and case names.
    
    Args:
        content: Direct text content
        file_path: Path to a file
        url: URL to fetch content from
        extract_case_names: Whether to extract case names
        
    Returns:
        Dictionary containing processing results
    """
    try:
        # Determine input type and extract text
        if content:
            logger.info("Processing direct text content")
            text = content
            source_type = "text"
            source_name = "pasted_text"
        elif file_path:
            logger.info(f"Processing file: {file_path}")
            text = extract_text_from_file(file_path)
            source_type = "file"
            source_name = Path(file_path).name
        elif url:
            logger.info(f"Processing URL: {url}")
            text = extract_text_from_url(url)
            source_type = "url"
            source_name = urlparse(url).netloc
        else:
            raise ValueError("Must provide content, file_path, or url")
        
        if not text.strip():
            logger.error("No text content extracted from input")
            return {
                "success": False,
                "error": "No text content extracted",
                "citations": [],
                "case_names": []
            }
        
        logger.info(f"Extracted text length: {len(text)} characters")
        logger.info(f"Sample of extracted text (first 500 characters): {text[:500]}...")
        
        # Extract citations and case names using the enhanced extraction logic
        try:
            # Use the enhanced citation extractor with case name and date extraction
            # DEPRECATED: from src.citation_extractor import CitationExtractor
            from src.enhanced_multi_source_verifier import EnhancedMultiSourceVerifier
            
            # Extract citations with case names and dates
            extractor = CitationExtractor(
                use_eyecite=True,
                use_regex=True,
                context_window=1000,
                deduplicate=True,
                extract_case_names=True
            )
            
            # Extract citations from text
            logger.info(f"[DEBUG] Starting citation extraction from text of length {len(text)}")
            logger.info(f"[DEBUG] Sample text: {text[:500]}...")
            
            extracted_citations = extractor.extract(text)
            
            # DEBUG: Log all extracted citations before any processing
            logger.info(f"[DEBUG] Raw extraction returned {len(extracted_citations)} citations")
            for i, citation in enumerate(extracted_citations, 1):
                logger.info(f"[DEBUG] Raw citation {i}: {citation}")
            
            # DEBUG: Check for specific citations we're looking for
            target_citations = [
                "200 Wn.2d 72", "514 P.3d 643",  # Convoyant
                "171 Wn.2d 486", "256 P.3d 321",  # Carlsen
                "146 Wn.2d 1", "43 P.3d 4"        # Campbell & Gwinn
            ]
            
            found_targets = []
            for citation in extracted_citations:
                citation_text = citation.get('citation', citation.get('citation_text', ''))
                for target in target_citations:
                    if target in citation_text:
                        found_targets.append(target)
                        logger.info(f"[DEBUG] Found target citation: {target} in {citation_text}")
            
            missing_targets = [t for t in target_citations if t not in found_targets]
            if missing_targets:
                logger.warning(f"[DEBUG] Missing target citations: {missing_targets}")
            else:
                logger.info(f"[DEBUG] All target citations found: {found_targets}")
            
            # Use the new simplified verification approach with fallback logic
            verified_citations = verify_citations_with_fallback(extracted_citations, text)
            
            # Convert to the expected format
            result = {
                'citations': verified_citations,
                'case_names': [],
                'metadata': {
                    'text_length': len(text),
                    'citations_found': len(verified_citations),
                    'extraction_method': 'enhanced_extraction_with_verification'
                }
            }
            
            initial_citation_count = len(result.get('citations', []))
            logger.info(f"Initially extracted {initial_citation_count} citations before verification")
            if initial_citation_count > 0:
                logger.info(f"All initial citations extracted before verification:")
                for i, citation in enumerate(result.get('citations', []), 1):
                    logger.info(f"Citation {i}: {citation.get('citation', 'N/A')}")
            # DEBUG: Log all extracted citations before verification
            logger.info(f"[DEBUG] Extracted {initial_citation_count} citations before verification")
            for i, citation in enumerate(result.get('citations', []), 1):
                logger.info(f"[DEBUG] Citation {i}: {citation.get('citation', 'N/A')}")
                
        except Exception as e:
            logger.error(f"Error in citation extraction: {str(e)}")
            # Fallback to empty result
            result = {
                'citations': [],
                'case_names': [],
                'metadata': {
                    'text_length': len(text),
                    'citations_found': 0,
                    'extraction_method': 'fallback',
                    'error': str(e)
                }
            }
            initial_citation_count = 0
        
        # Extract case names from citations if they have case_name fields
        case_names = []
        if extract_case_names and result.get("citations"):
            for citation in result.get("citations", []):
                case_name = citation.get("case_name", "")
                if case_name and case_name != "N/A" and case_name not in case_names:
                    case_names.append(case_name)
            logger.info(f"Extracted {len(case_names)} unique case names from citations")
        if extract_case_names and not case_names:
            extractor = CitationExtractor()
            case_names = extractor._extract_case_names_from_text(text)
            logger.info(f"Extracted {len(case_names)} unique case names directly from text")
        # DEBUG: Log the full result returned
        logger.info(f"[DEBUG] Returning from process_document: success={True}, citations={len(result.get('citations', []))}, case_names={len(case_names)}")
        return {
            "success": True,
            "text_length": len(text),
            "source_type": source_type,
            "source_name": source_name,
            "citations": result.get("citations", []),
            "case_names": case_names,
            "extraction_metadata": result.get("metadata", {}),
            "statistics": result.get("statistics", {}),
            "summary": result.get("summary", {}),
            "verification_metadata": result.get("verification_metadata", {})
        }
        
    except Exception as e:
        logger.error(f"Error processing document: {str(e)}")
        return {
            "success": False,
            "error": str(e),
            "citations": [],
            "case_names": []
        }


def process_document_input(input_type, input_value, file_ext=None):
    """
    Unified document input handler for file, url, or text.
    Returns: (text, metadata)
    """
    logger = logging.getLogger(__name__)
    meta = {}
    text = None
    try:
        if input_type == 'file':
            from src.file_utils import extract_text_from_file
            text_result = extract_text_from_file(input_value, file_ext=file_ext)
            if isinstance(text_result, tuple):
                text, extracted_case_name = text_result
                meta['extracted_case_name'] = extracted_case_name
            else:
                text = text_result
            meta['source_type'] = 'file'
            meta['source_name'] = input_value
        elif input_type == 'url':
            from src.enhanced_validator_production import extract_text_from_url
            text_result = extract_text_from_url(input_value)
            text = text_result.get('text', '')
            meta['content_type'] = text_result.get('content_type')
            meta['source_type'] = 'url'
            meta['source_name'] = input_value
            meta['status'] = text_result.get('status')
            meta['error'] = text_result.get('error')
        elif input_type == 'text':
            text = input_value
            meta['source_type'] = 'text'
            meta['source_name'] = 'pasted_text'
        else:
            raise ValueError("Unknown input type")
        meta['text_length'] = len(text) if text else 0
    except Exception as e:
        logger.error(f"[process_document_input] Error processing {input_type}: {e}")
        meta['error'] = str(e)
        text = ''
    return text, meta


def extract_and_verify_citations(text, use_enhanced=True, logger=None):
    """
    Extract and verify citations from text with performance optimizations.
    
    Args:
        text: Text to extract citations from
        use_enhanced: Whether to use enhanced extraction
        logger: Logger instance
        
    Returns:
        List of citation results with verification
    """
    if not logger:
        logger = logging.getLogger(__name__)
    
    start_time = time.time()
    
    # Extract citations using UnifiedCitationProcessor (not deprecated CitationExtractor)
    logger.info("🔍 [ANALYZE] Extracting citations using UnifiedCitationProcessor...")
    from src.unified_citation_processor import unified_processor
    
    # Process text with unified processor
    result = unified_processor.process_text(text, options={
        'extract_case_names': True,
        'use_enhanced': use_enhanced
    })
    
    # The unified processor returns {'results': [...], 'statistics': ..., 'summary': ...}
    # We need to use 'results' not 'citations'
    citations = result.get('results', [])
    extraction_time = time.time() - start_time
    
    logger.info(f"✅ [ANALYZE] Extracted {len(citations)} citations in {extraction_time:.2f}s")
    
    if not citations:
        return []
    
    # The unified processor already handles verification, so we just need to format for frontend
    logger.info("🔍 [ANALYZE] Formatting results for frontend...")
    formatting_start = time.time()
    
    # Format citations for frontend compatibility
    formatted_citations = []
    for citation in citations:
        # Convert CitationResult to dict format expected by frontend
        formatted_citation = {
            'citation': citation.citation,
            'verified': 'true' if citation.verified else 'false',
            'case_name': citation.case_name or 'N/A',
            'extracted_case_name': citation.extracted_case_name or 'N/A',
            'canonical_name': citation.canonical_name or 'N/A',
            'canonical_date': citation.canonical_date or 'N/A',
            'extracted_date': citation.extracted_date or 'N/A',
            'court': citation.court or 'N/A',
            'docket_number': citation.docket_number or 'N/A',
            'confidence': citation.confidence,
            'source': citation.source,
            'url': citation.url or '',
            'context': citation.context or '',
            'is_complex_citation': citation.is_complex,
            'is_parallel_citation': citation.is_parallel,
            'error': citation.error or '',
            'method': citation.method,
            'pattern': citation.pattern
        }
        # DEBUG: Log extracted fields
        logger.info(f"[DEBUG] Citation: {citation.citation} | Extracted Name: {citation.extracted_case_name} | Extracted Date: {citation.extracted_date}")
        
        # Format parallel citations as objects with verification status (not simple strings)
        if citation.parallel_citations:
            formatted_citation['parallels'] = []
            for parallel_citation in citation.parallel_citations:
                # For now, assume parallel citations are verified if the main citation is verified
                # In a more sophisticated implementation, each parallel would be verified individually
                parallel_obj = {
                    'citation': parallel_citation,
                    'verified': citation.verified,  # Use main citation's verification status
                    'true_by_parallel': False,  # This would be set if verified by another parallel
                    'case_name': citation.case_name,
                    'extracted_case_name': citation.extracted_case_name,
                    'extracted_date': citation.extracted_date
                }
                formatted_citation['parallels'].append(parallel_obj)
        else:
            formatted_citation['parallels'] = []
        
        # Keep the old parallel_citations field for backward compatibility
        formatted_citation['parallel_citations'] = citation.parallel_citations or []
        
        formatted_citations.append(formatted_citation)
    
    formatting_time = time.time() - formatting_start
    total_time = time.time() - start_time
    
    logger.info(f"✅ [ANALYZE] Formatting completed in {formatting_time:.2f}s")
    logger.info(f"✅ [ANALYZE] Total processing time: {total_time:.2f}s")
    logger.info(f"✅ [ANALYZE] Returning response with {len(formatted_citations)} citations")
    
    return formatted_citations


def verify_citations_with_fallback(citations: List[Dict], text: str) -> List[Dict]:
    """
    Verify all citations individually, then apply fallback logic for parallel citations.
    
    This simplified approach:
    1. Verifies every citation individually 
    2. Groups citations by parallel_group_id for fallback logic
    3. If a citation fails but one of its parallels succeeds, marks it as verified_by_parallel
    
    Args:
        citations: List of citation dictionaries from extraction
        text: Original text for context
        
    Returns:
        List of verified citations with fallback logic applied
    """
    from src.enhanced_multi_source_verifier import EnhancedMultiSourceVerifier
    
    logger.info(f"[VERIFY] Starting verification of {len(citations)} citations")
    
    # Initialize verifier
    verifier = EnhancedMultiSourceVerifier()
    
    # Step 1: Verify every citation individually (skip combined citations)
    verified_citations = []
    for citation_data in citations:
        citation_text = citation_data.get('citation', citation_data.get('citation_text', ''))
        if not citation_text:
            continue
        # Skip combined citations (they have is_parallel_group=True)
        if citation_data.get('is_parallel_group', False):
            logger.info(f"[VERIFY] Skipping combined citation for verification: {citation_text}")
            continue
        # Log the citation string being sent to the verifier
        logger.info(f"[VERIFY] Sending citation to verifier: '{citation_text}' (raw)")
        # If you have normalization, log that too
        from src.citation_extractor import normalize_washington_citations
        normalized_citation = normalize_washington_citations(citation_text)
        logger.info(f"[VERIFY] Normalized citation: '{normalized_citation}'")
        # Get extracted case name and date from the citation data
        extracted_case_name = citation_data.get('case_name', '')
        extracted_date = citation_data.get('date', '')
        # Verify citation with extracted information
        verification_result = verifier.verify_citation_unified_workflow(
            normalized_citation,
            extracted_case_name=extracted_case_name,
            extracted_date=extracted_date
        )
        # Merge verification result with original citation data
        merged_result = {
            **citation_data,
            'verified': verification_result.get('verified', False),
            'url': verification_result.get('url', ''),
            'court': verification_result.get('court', ''),
            'docket_number': verification_result.get('docket_number', ''),
            'canonical_date': verification_result.get('canonical_date', ''),
            'source': verification_result.get('source', 'Unknown'),
            'confidence': verification_result.get('confidence', 0.0),
            'error': verification_result.get('error', ''),
            'canonical_name': verification_result.get('canonical_name', ''),
            'extracted_case_name': extracted_case_name,
            'extracted_date': extracted_date,
            'context': citation_data.get('context', text),
            'verified_by_parallel': False  # Will be set in step 2
        }
        verified_citations.append(merged_result)
    
    # Step 2: Apply fallback logic for parallel citations
    # Group citations by parallel_group_id
    parallel_groups = {}
    for citation in verified_citations:
        group_id = citation.get('parallel_group_id')
        if group_id is not None:
            if group_id not in parallel_groups:
                parallel_groups[group_id] = []
            parallel_groups[group_id].append(citation)
    
    # Apply fallback logic: if any citation in a group is verified, mark unverified ones as verified_by_parallel
    for group_id, group_citations in parallel_groups.items():
        # Check if any citation in the group is verified
        any_verified = any(citation.get('verified', False) for citation in group_citations)
        
        if any_verified:
            # Mark unverified citations in the group as verified_by_parallel
            for citation in group_citations:
                if not citation.get('verified', False):
                    citation['verified_by_parallel'] = True
                    citation['verified'] = True  # Also set verified to true for UI consistency
                    citation['source'] = 'verified_by_parallel'
                    logger.info(f"[VERIFY] Citation '{citation.get('citation')}' marked as verified_by_parallel")
    
    # Step 3: Add back combined citations for display purposes
    for citation_data in citations:
        if citation_data.get('is_parallel_group', False):
            # Find the components of this combined citation
            group_id = citation_data.get('parallel_group_id')
            if group_id is not None:
                # Check if any component is verified
                component_verified = any(
                    citation.get('verified', False) 
                    for citation in verified_citations 
                    if citation.get('parallel_group_id') == group_id
                )
                
                # Create combined citation result
                combined_result = {
                    **citation_data,
                    'verified': component_verified,
                    'source': 'combined_display',
                    'verified_by_parallel': component_verified,
                    'context': citation_data.get('context', text)
                }
                verified_citations.append(combined_result)
    
    logger.info(f"[VERIFY] Completed verification with fallback logic. {len(verified_citations)} citations processed.")
    return verified_citations 