"""
Citation utilities for CaseStrainer.

This module provides functions for citation extraction, validation, and processing
using the unified citation processor system.
"""

import re
import logging
import time
from typing import List, Dict, Any, Optional, Tuple, Set
from dataclasses import dataclass, asdict
import unicodedata
import warnings
from collections import defaultdict, deque
import os
import sys
from datetime import datetime
import json
import hashlib
import sqlite3
from pathlib import Path

# Add src to path for imports
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

# Import citation utilities from consolidated module
try:
    from src.citation_utils_consolidated import apply_washington_spacing_rules
except ImportError:
    from citation_utils_consolidated import apply_washington_spacing_rules

# Add the project root to the Python path
project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if project_root not in sys.path:
    sys.path.insert(0, project_root)

# Import from config with absolute path
from config import ALLOWED_EXTENSIONS

# Setup logger (modules importing this should configure logging)
logger = logging.getLogger(__name__)


def get_unified_citations(text: str, logger: Optional[logging.Logger] = None) -> List[str]:
    """Get citations using the new unified processor with eyecite."""
    from unified_citation_processor_v2 import UnifiedCitationProcessorV2, ProcessingConfig
    
    config = ProcessingConfig(
        use_eyecite=True,
        use_regex=True,
        extract_case_names=True,
        extract_dates=True,
        enable_clustering=True,
        enable_deduplication=True,
        debug_mode=False
    )
    
    processor = UnifiedCitationProcessorV2(config)
    results = processor.process_text(text)
    
    # Return just the citation strings for compatibility
    return [result.citation for result in results]


def allowed_file(filename: str) -> bool:
    """Check if a file has an allowed extension."""
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def log_citation_verification(citation: str, result: Dict[str, Any]) -> None:
    """
    Log the result of a citation verification attempt.
    
    Args:
        citation: The citation text that was verified.
        result: The result of the verification process.
    """
    status = "SUCCESS" if result.get("verified", False) else "FAILED"
    logger.info(f"Citation Verification [{status}]: {citation}")
    if result.get("error", False):
        logger.error(f"Verification error: {result.get('error', 'Unknown error')}")
    elif result.get("verified", False):
        try:
            logger.info(f"Verified citation details: {result.get('source', 'No source info')}")
        except UnicodeEncodeError:
            safe_source = str(result.get('source', 'No source info')).encode('cp1252', errors='replace').decode('cp1252')
            logger.info(f"Verified citation details (safe): {safe_source}")


def normalize_citation_text(citation_text: str) -> str:
    """
    Normalize citation text to a standard format before processing.

    Handles common issues like:
    - Extra spaces in reporter abbreviations (e.g., 'F. 3d' -> 'F.3d')
    - Double periods (e.g., 'U.S..' -> 'U.S.')
    - Inconsistent spacing around v. (e.g., 'U.S. v.Caraway' -> 'U.S. v. Caraway')
    - Washington citation spacing rules (Wn.2d vs Wn. App.)

    Args:
        citation_text: The citation text to normalize

    Returns:
        The normalized citation text
    """
    if not citation_text or not isinstance(citation_text, str):
        return citation_text

    # Remove any leading/trailing whitespace
    normalized = citation_text.strip()

    # Fix double periods
    normalized = re.sub(r"\.\.+", ".", normalized)

    # Apply Washington citation spacing rules FIRST (before general reporter fixes)
    normalized = apply_washington_spacing_rules(normalized)

    # Fix spaces in reporter abbreviations (e.g., 'F. 3d' -> 'F.3d')
    # But exclude Washington citations that we just fixed
    normalized = re.sub(r"(\b[A-Za-z]+\.)\s+(\d+[a-z]*)(?!\s*Wn\.)", r"\1\2", normalized)

    # Fix spacing around 'v.'
    normalized = re.sub(r"\s+v\.\s*", " v. ", normalized, flags=re.IGNORECASE)

    # Fix common reporter abbreviations
    reporter_fixes = {
        r"\bFed\.\s*": "F.",
        r"\bFed\.\s*App\.\s*": "F. App'x",
        r"\bF\.\s*2d\b": "F.2d",
        r"\bF\.\s*3d\b": "F.3d",
        r"\bF\.\s*4th\b": "F.4th",
        r"\bU\.\s*S\.\s*": "U.S. ",
        r"\bS\.\s*Ct\.\s*": "S. Ct. ",
        r"\bL\.\s*Ed\.\s*2d\b": "L. Ed. 2d",
    }

    for pattern, replacement in reporter_fixes.items():
        normalized = re.sub(pattern, replacement, normalized)

    # Remove extra spaces
    normalized = " ".join(normalized.split())

    return normalized


def verify_citation(
    citation: str,
    context: Optional[str] = None,
    logger: Optional[logging.Logger] = None,
    DEFAULT_API_KEY: Optional[str] = None,
    thread_local: Optional[Any] = None,
    timeout: int = 15,
    extracted_case_name: Optional[str] = None
) -> Dict[str, Any]:
    """Verify a citation using the unified citation processor."""
    try:
        from unified_citation_processor_v2 import UnifiedCitationProcessorV2 as UnifiedCitationProcessor
        processor = UnifiedCitationProcessor()
        return processor.verify_citation(citation, extracted_case_name=extracted_case_name)
    except ImportError:
        # Fallback to basic verification
        if logger:
            logger.warning("Unified citation processor not available, using fallback")
        return {
            "citation": citation,
            "verified": False,
            "error": "Citation processor not available"
        }


def extract_citations_from_file(filepath: str, logger: Optional[logging.Logger] = None) -> List[Dict[str, Any]]:
    """Extract citations from a file and return full metadata."""
    if logger is None:
        logger = logging.getLogger(__name__)
        
    try:
        logger.info(f"[DEBUG] Starting extract_citations_from_file for: {filepath}")
        file_size = os.path.getsize(filepath)
        file_extension = os.path.splitext(filepath)[1].lower()
        logger.info(
            f"[DEBUG] File details - Size: {file_size} bytes, Extension: {file_extension}"
        )
        start_time = time.time()
        text = ""
        
        if filepath.endswith(".pdf"):
            logger.info("Detected PDF file, using PDFHandler for extraction")
            # Always use PDFMiner as the only extraction method
            from pdf_handler import PDFHandler, PDFExtractionConfig, PDFExtractionMethod
            handler = PDFHandler(PDFExtractionConfig(
                preferred_method=PDFExtractionMethod.PDFMINER,
                use_fallback=False,  # Only use PDFMiner
                timeout=30,  # 30 second timeout
                clean_text=True
            ))
            text = handler.extract_text(filepath)
            if text.startswith("Error:"):
                logger.error(f"PDF extraction failed: {text}")
                raise Exception(text)
            logger.info(f"Successfully extracted {len(text)} characters from PDF")
            
        elif filepath.endswith((".doc", ".docx")):
            logger.info("Detected Word document, using python-docx for extraction")
            import docx

            try:
                doc = docx.Document(filepath)
                logger.info(f"Word document has {len(doc.paragraphs)} paragraphs")
                paragraphs = [paragraph.text for paragraph in doc.paragraphs]
                text = "\n".join(paragraphs)
                logger.info(
                    f"Extracted {len(paragraphs)} paragraphs with total {len(text)} characters"
                )
            except Exception as docx_error:
                logger.error(f"Error reading Word document: {str(docx_error)}")
                raise
        else:
            logger.info(f"Using standard text file reading for {file_extension} file")
            try:
                with open(filepath, "r", encoding="utf-8") as file:
                    text = file.read()
                logger.info(f"Successfully read {len(text)} characters from text file")
            except UnicodeDecodeError:
                logger.warning("UTF-8 decode failed, trying with error handling")
                with open(filepath, "r", encoding="utf-8", errors="replace") as file:
                    text = file.read()
                logger.info(f"Successfully read {len(text)} characters with error handling")
            except Exception as e:
                logger.error(f"Error reading text file: {str(e)}")
                raise

        # Extract citations from the text
        citations = extract_citations_from_text(text, logger)
        
        processing_time = time.time() - start_time
        logger.info(f"Citation extraction completed in {processing_time:.2f} seconds")
        logger.info(f"Extracted {len(citations)} citations from file")
        
        return citations
        
    except Exception as e:
        logger.error(f"Error in extract_citations_from_file: {str(e)}")
        logger.error(traceback.format_exc())
        raise


def extract_citations_from_text(text: str, logger: Optional[logging.Logger] = None) -> List[Dict[str, Any]]:
    """Extract citations from text using the unified processor."""
    if logger is None:
        logger = logging.getLogger(__name__)
        
    if not text or not isinstance(text, str):
        logger.warning("Invalid text input for citation extraction")
        return []

    try:
        logger.info(f"Starting citation extraction from text ({len(text)} characters)")
        
        # Use the unified citation processor
        citation_strings = get_unified_citations(text, logger)
        
        # Convert to the expected format
        citations = []
        for citation_str in citation_strings:
            citations.append({
                "citation": citation_str,
                "extracted_from": "text",
                "confidence": 0.8,  # Default confidence for extracted citations
                "verified": False,
                "source": "extraction"
            })
        
        logger.info(f"Extracted {len(citations)} citations from text")
        return citations
        
    except Exception as e:
        logger.error(f"Error extracting citations from text: {str(e)}")
        logger.error(traceback.format_exc())
        return []


def get_citation_context(text: str, citation_text: str, context_size: int = 250) -> str:
    """Get context around a citation in the text."""
    if not text or not citation_text:
        return ""
    
    try:
        # Find the citation in the text
        index = text.find(citation_text)
        if index == -1:
            return ""
        
        # Calculate start and end positions for context
        start = max(0, index - context_size)
        end = min(len(text), index + len(citation_text) + context_size)
        
        # Extract context
        context = text[start:end]
        
        # Add ellipsis if we're not at the beginning/end
        if start > 0:
            context = "..." + context
        if end < len(text):
            context = context + "..."
        
        return context
        
    except Exception as e:
        logger.error(f"Error getting citation context: {str(e)}")
        return ""


def validate_potential_citation(citation_text: str) -> bool:
    """Validate if a string looks like a potential citation."""
    if not citation_text or not isinstance(citation_text, str):
        return False
    
    # Basic citation patterns
    citation_patterns = [
        r'\d+\s+[A-Za-z\.]+\s+\d+',  # Volume Reporter Page
        r'\d+\s+[A-Za-z\.]+\s*\(\d{4}\)',  # Volume Reporter (Year)
        r'\d+\s+U\.S\.\s+\d+',  # US Supreme Court
        r'\d+\s+F\.(?:2d|3d|4th)?\s+\d+',  # Federal Reporter
        r'\d+\s+[A-Za-z\.]+\s+(?:2d|3d|4th|5th|6th|7th|8th|9th)\s+\d+',  # Series reporters
    ]
    
    for pattern in citation_patterns:
        if re.search(pattern, citation_text, re.IGNORECASE):
            return True
    
    return False


def clean_and_validate_citations(citations: List[str]) -> Tuple[List[str], Dict[str, int]]:
    """Clean and validate a list of citations."""
    if not citations:
        return [], {"total_removed": 0, "empty_strings": 0, "malformed": 0, "secondary_validated": 0}
    
    cleaned_citations = []
    stats = {
        "total_removed": 0,
        "empty_strings": 0,
        "malformed": 0,
        "secondary_validated": 0
    }
    
    for citation in citations:
        if not citation or not isinstance(citation, str):
            stats["empty_strings"] += 1
            continue
        
        citation = citation.strip()
        if not citation:
            stats["empty_strings"] += 1
            continue
        
        # Basic validation
        if not validate_potential_citation(citation):
            stats["malformed"] += 1
            continue
        
        # Normalize the citation
        normalized = normalize_citation_text(citation)
        if normalized and normalized not in cleaned_citations:
            cleaned_citations.append(normalized)
            stats["secondary_validated"] += 1
    
    stats["total_removed"] = len(citations) - len(cleaned_citations)
    
    return cleaned_citations, stats


def batch_validate_citations_optimized(citations: List[str], api_key: Optional[str] = None) -> List[Dict[str, Any]]:
    """Optimized batch validation of citations."""
    if not citations:
        return []
    
    # Clean and validate citations before processing
    cleaned_citations, validation_stats = clean_and_validate_citations(citations)
    
    if validation_stats["total_removed"] > 0:
        logger.warning(
            f"Removed {validation_stats['total_removed']} invalid citations before processing: "
            f"{validation_stats['empty_strings']} empty strings, "
            f"{validation_stats['malformed']} malformed citations. "
            f"Additionally, {validation_stats['secondary_validated']} citations passed secondary validation."
        )
    
    if not cleaned_citations:
        return []

    results = []
    
    def validate_single_citation(citation: str) -> Dict[str, Any]:
        """Validate a single citation."""
        try:
            from unified_citation_processor_v2 import UnifiedCitationProcessorV2 as UnifiedCitationProcessor
            verifier = UnifiedCitationProcessor()
            result = verifier.verify_citation_unified_workflow(citation)
            return {
                "citation": citation,
                "exists": result.get("verified", False),
                "method": result.get("verification_method", "unknown"),
                "error": None,
                "data": result,
            }
        except Exception as e:
            logger.error(f"Error validating citation {citation}: {str(e)}")
            return {
                "citation": citation,
                "exists": False,
                "method": "error",
                "error": str(e),
                "data": None,
            }
    
    # Process citations in batches
    batch_size = 10
    for i in range(0, len(cleaned_citations), batch_size):
        batch = cleaned_citations[i:i + batch_size]
        batch_results = [validate_single_citation(citation) for citation in batch]
        results.extend(batch_results)
    
    return results


def batch_validate_citations(citations: List[str], api_key: Optional[str] = None) -> List[Dict[str, Any]]:
    """Batch validation of citations (legacy function)."""
    return batch_validate_citations_optimized(citations, api_key)


# DEPRECATED: Use extract_all_citations from .citation_extractor instead.
def extract_citations(*args: Any, **kwargs: Any) -> List[Any]:
    """DEPRECATED. Use extract_all_citations from .citation_extractor instead."""
    import warnings
    warnings.warn("extract_citations is deprecated. Use extract_all_citations from .citation_extractor instead.", DeprecationWarning)
    return []


def chunk_text_with_overlap(text: str, max_len: int = 64000, overlap: int = 40) -> List[str]:
    """Split text into chunks of max_len with specified overlap."""
    chunks = []
    i = 0
    while i < len(text):
        start = i - overlap if i > 0 else 0
        end = min(i + max_len, len(text))
        chunks.append(text[start:end])
        i += max_len
    return chunks


def deduplicate_citations(citations: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Remove duplicate citations from a list."""
    seen = set()
    unique = []
    for c in citations:
        key = c.get('normalized_citations', [None])[0] if c.get('normalized_citations') else c.get('citation') or c.get('citation_text')
        if key and key not in seen:
            seen.add(key)
            unique.append(c)
    return unique
