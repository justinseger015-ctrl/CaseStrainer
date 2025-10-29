#!/usr/bin/env python3
"""
Unified Text Extraction for All File Formats
Fast and reliable text extraction from multiple document formats.

Enhanced with multi-method PDF extraction, advanced text normalization,
and content caching for improved performance.
"""

import logging
import re
import hashlib
import time
from pathlib import Path
from typing import Optional, Tuple, Dict, Any
import tempfile
import os

logger = logging.getLogger(__name__)

# Extract timeout per format (seconds)
EXTRACTION_TIMEOUT = 30

# Cache for processed content
_content_cache = {}
_cache_ttl = 3600  # 1 hour


class UnifiedTextExtractor:
    """
    Unified text extraction for all supported formats.
    
    Supported formats:
    - PDF: via PyMuPDF, PyPDF2, and robust_pdf_extractor with fallbacks
    - DOCX: via python-docx
    - DOC: via antiword or textract
    - HTML/HTM/XML: via BeautifulSoup
    - RTF: via striprtf
    - TXT/MD: direct read with enhanced normalization
    
    Enhanced features:
    - Multi-method PDF extraction with automatic fallback
    - Advanced text normalization and artifact removal
    - Content caching for performance optimization
    """
    
    def __init__(self, verbose: bool = False):
        """Initialize extractor with format detection and enhanced capabilities."""
        self.verbose = verbose
        self._check_dependencies()
        self._init_pdf_extractors()
    
    def _check_dependencies(self):
        """Check which libraries are available."""
        self.has_docx = self._check_import('docx')
        self.has_bs4 = self._check_import('bs4')
        self.has_striprtf = self._check_import('striprtf')
        
        if self.verbose:
            logger.info(f"Text extraction capabilities: DOCX={self.has_docx}, HTML={self.has_bs4}, RTF={self.has_striprtf}")
    
    def _check_import(self, module_name: str) -> bool:
        """Check if a module is available."""
        try:
            __import__(module_name)
            return True
        except ImportError:
            return False
    
    def _init_pdf_extractors(self):
        """Initialize PDF extraction capabilities."""
        self.has_pymupdf = self._check_import('fitz')
        self.has_pypdf2 = self._check_import('PyPDF2')
        self.has_robust_extractor = self._check_import('robust_pdf_extractor')
        
        if self.verbose:
            logger.info(f"PDF extraction capabilities: PyMuPDF={self.has_pymupdf}, PyPDF2={self.has_pypdf2}, Robust={self.has_robust_extractor}")
    
    def _get_cache_key(self, file_path: str) -> str:
        """Generate cache key for file."""
        try:
            # Use file path, size, and modification time for cache key
            stat = os.stat(file_path)
            cache_data = f"{file_path}_{stat.st_size}_{stat.st_mtime}"
            return hashlib.md5(cache_data.encode('utf-8')).hexdigest()
        except Exception:
            # Fallback to path hash only
            return hashlib.md5(file_path.encode('utf-8')).hexdigest()
    
    def _get_from_cache(self, cache_key: str) -> Optional[str]:
        """Get text from cache if available and not expired."""
        if cache_key in _content_cache:
            cached_data = _content_cache[cache_key]
            if time.time() - cached_data['timestamp'] < _cache_ttl:
                if self.verbose:
                    logger.debug(f"Cache hit for {cache_key}")
                return cached_data['text']
            else:
                # Remove expired cache entry
                del _content_cache[cache_key]
        return None
    
    def _cache_content(self, cache_key: str, text: str):
        """Cache extracted text."""
        _content_cache[cache_key] = {
            'text': text,
            'timestamp': time.time()
        }
        if self.verbose:
            logger.debug(f"Cached content for {cache_key}")
    
    def _enhanced_text_normalization(self, text: str) -> str:
        """Enhanced text normalization with multiple cleaning steps."""
        if not text:
            return text
        
        try:
            # Step 1: Basic Unicode normalization
            from utils.text_normalizer import normalize_text
            normalized = normalize_text(text)
            
            # Step 2: Remove problematic characters
            normalized = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', normalized)
            
            # Step 3: Normalize whitespace
            normalized = re.sub(r'\r\n', '\n', normalized)  # Windows newlines
            normalized = re.sub(r'\r', '\n', normalized)    # Old Mac newlines
            normalized = re.sub(r'\n{3,}', '\n\n', normalized)  # Multiple newlines
            normalized = re.sub(r'[ \t]+', ' ', normalized)  # Multiple spaces/tabs
            normalized = re.sub(r' *\n *', '\n', normalized)  # Spaces around newlines
            
            # Step 4: Remove common document artifacts
            normalized = re.sub(r'\n\d+\n', '\n', normalized)  # Page numbers
            normalized = re.sub(r'\n-+\s*\d+\s*-+\n', '\n', normalized)  # Header/footer
            normalized = re.sub(r'\S+@\S+\.\S+', '', normalized)  # Email addresses
            
            # Step 5: Final cleanup
            normalized = normalized.strip()
            
            return normalized
            
        except Exception as e:
            logger.warning(f"Error in text normalization: {str(e)}")
            return text  # Return original if normalization fails
    
    def extract_text(self, file_path: str) -> Tuple[str, str]:
        """
        Extract text from any supported file format with enhanced processing.
        
        Args:
            file_path: Path to file
            
        Returns:
            Tuple of (extracted_text, method_used)
        """
        if not Path(file_path).exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Check cache first
        cache_key = self._get_cache_key(file_path)
        cached_text = self._get_from_cache(cache_key)
        if cached_text:
            return cached_text, "cache"
        
        start_time = time.time()
        file_path_obj = Path(file_path)
        ext = file_path_obj.suffix.lower().lstrip('.')
        
        try:
            # Route to appropriate extractor
            if ext == 'pdf':
                text, method = self._extract_pdf_enhanced(file_path)
            elif ext in ['docx', 'doc']:
                text, method = self._extract_word(file_path, ext)
            elif ext in ['html', 'htm', 'xml', 'xhtml']:
                text, method = self._extract_html(file_path)
            elif ext == 'rtf':
                text, method = self._extract_rtf(file_path)
            elif ext in ['txt', 'md', 'markdown']:
                text, method = self._extract_plaintext_enhanced(file_path)
            else:
                # Fallback: try as plain text
                text, method = self._extract_plaintext_enhanced(file_path)
                method = f"{ext}:fallback_plaintext"
            
            # Apply enhanced text normalization
            if text:
                text = self._enhanced_text_normalization(text)
            
            elapsed = time.time() - start_time
            
            # Cache successful extractions
            if text and len(text.strip()) > 50:
                self._cache_content(cache_key, text)
                logger.info(f"✅ Extracted {len(text):,} chars from {ext} in {elapsed:.1f}s using {method}")
                return text, method
            else:
                logger.warning(f"⚠️ Insufficient text extracted from {ext}: {len(text) if text else 0} chars")
                return "", f"{ext}:insufficient_text"
                
        except Exception as e:
            elapsed = time.time() - start_time
            logger.error(f"❌ Failed to extract {ext} after {elapsed:.1f}s: {e}")
            return "", f"{ext}:error"
    
    def _extract_pdf_enhanced(self, file_path: str) -> Tuple[str, str]:
        """Extract text from PDF using multiple methods with fallbacks."""
        methods_tried = []
        
        # Method 1: PyMuPDF (fitz) - Preferred for quality and speed
        if self.has_pymupdf:
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(file_path)
                text_parts = []
                
                for page_num in range(min(5, len(doc))):  # Test first 5 pages
                    page = doc[page_num]
                    
                    # Get page dimensions
                    rect = page.rect
                    width, height = rect.width, rect.height
                    
                    # Define text area (exclude headers and footers)
                    text_area = fitz.Rect(0, 65, width, height - 50)
                    
                    # Extract text from defined area
                    text = page.get_text("text", clip=text_area)
                    
                    if text.strip():
                        text_parts.append(text)
                
                doc.close()
                pymupdf_text = "\n".join(text_parts)
                
                if pymupdf_text and len(pymupdf_text.strip()) > 100:
                    methods_tried.append("PyMuPDF")
                    return pymupdf_text, "pdf:pymupdf"
                    
            except Exception as e:
                if self.verbose:
                    logger.warning(f"PyMuPDF extraction failed: {e}")
                methods_tried.append("PyMuPDF(failed)")
        
        # Method 2: PyPDF2 - Reliable fallback
        if self.has_pypdf2:
            try:
                import PyPDF2
                
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    text_parts = []
                    
                    for page in reader.pages[:5]:  # Test first 5 pages
                        text = page.extract_text()
                        if text.strip():
                            text_parts.append(text)
                
                pypdf2_text = "\n".join(text_parts)
                
                if pypdf2_text and len(pypdf2_text.strip()) > 100:
                    methods_tried.append("PyPDF2")
                    return pypdf2_text, "pdf:pypdf2"
                    
            except Exception as e:
                if self.verbose:
                    logger.warning(f"PyPDF2 extraction failed: {e}")
                methods_tried.append("PyPDF2(failed)")
        
        # Method 3: RobustPDFExtractor - Final fallback
        if self.has_robust_extractor:
            try:
                from src.robust_pdf_extractor import extract_pdf_text_robust
                text, library = extract_pdf_text_robust(file_path, verbose=self.verbose)
                
                if text and len(text.strip()) > 100:
                    methods_tried.append(f"Robust({library})")
                    return text, f"pdf:robust_{library}"
                    
            except Exception as e:
                if self.verbose:
                    logger.warning(f"RobustPDFExtractor failed: {e}")
                methods_tried.append("Robust(failed)")
        
        logger.error(f"❌ All PDF extraction methods failed: {', '.join(methods_tried)}")
        return "", "pdf:all_methods_failed"
    
    def _extract_pdf(self, file_path: str) -> Tuple[str, str]:
        """Extract text from PDF using robust extractor (legacy method)."""
        from src.robust_pdf_extractor import extract_pdf_text_robust
        text, library = extract_pdf_text_robust(file_path, verbose=self.verbose)
        return text, f"pdf:{library}"
    
    def _extract_word(self, file_path: str, ext: str) -> Tuple[str, str]:
        """Extract text from Word documents (.docx or .doc)."""
        if ext == 'docx' and self.has_docx:
            return self._extract_docx(file_path)
        elif ext == 'doc':
            # .doc requires external tools - try multiple methods
            return self._extract_doc_legacy(file_path)
        else:
            # No docx library available
            logger.warning("python-docx not available, cannot extract .docx files")
            return "", "docx:not_available"
    
    def _extract_docx(self, file_path: str) -> Tuple[str, str]:
        """Extract text from .docx using python-docx."""
        try:
            import docx
            doc = docx.Document(file_path)
            
            # Extract all paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            text = '\n\n'.join(paragraphs)
            
            # Also extract tables if present
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text += '\n' + row_text
            
            return text, "docx:python-docx"
            
        except Exception as e:
            if self.verbose:
                logger.warning(f"python-docx extraction failed: {e}")
            return "", "docx:error"
    
    def _extract_doc_legacy(self, file_path: str) -> Tuple[str, str]:
        """
        Extract text from legacy .doc format.
        Tries multiple methods: antiword, textract, direct read.
        """
        # Method 1: Try antiword (if available)
        try:
            import subprocess
            result = subprocess.run(
                ['antiword', file_path],
                capture_output=True,
                text=True,
                timeout=EXTRACTION_TIMEOUT
            )
            if result.returncode == 0 and result.stdout:
                return result.stdout, "doc:antiword"
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass
        
        # Method 2: Try textract (if available)
        try:
            import textract
            text = textract.process(file_path, encoding='utf-8').decode('utf-8')
            if text:
                return text, "doc:textract"
        except ImportError:
            pass
        except Exception:
            pass
        
        # Method 3: Try reading as plain text (low quality)
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            if text and len(text.strip()) > 100:
                logger.warning("⚠️ Using low-quality plain text extraction for .doc file")
                return text, "doc:plaintext_fallback"
        except:
            pass
        
        logger.error("❌ No method available to extract .doc file. Install antiword or textract.")
        return "", "doc:no_extractor"
    
    def _extract_html(self, file_path: str) -> Tuple[str, str]:
        """Extract text from HTML/XML files."""
        if not self.has_bs4:
            logger.warning("BeautifulSoup not available, using plain text extraction")
            return self._extract_plaintext(file_path)
        
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for element in soup(['script', 'style', 'head', 'meta']):
                element.decompose()
            
            # Get text
            text = soup.get_text(separator=' ', strip=True)
            
            # Clean up whitespace
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            text = '\n'.join(lines)
            
            return text, "html:beautifulsoup"
            
        except Exception as e:
            if self.verbose:
                logger.warning(f"BeautifulSoup extraction failed: {e}, trying plain text")
            return self._extract_plaintext(file_path)
    
    def _extract_rtf(self, file_path: str) -> Tuple[str, str]:
        """Extract text from RTF files."""
        if not self.has_striprtf:
            logger.warning("striprtf not available, using plain text extraction")
            return self._extract_plaintext(file_path)
        
        try:
            from striprtf.striprtf import rtf_to_text
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                rtf_content = f.read()
            
            text = rtf_to_text(rtf_content)
            return text, "rtf:striprtf"
            
        except Exception as e:
            if self.verbose:
                logger.warning(f"striprtf extraction failed: {e}, trying plain text")
            return self._extract_plaintext(file_path)
    
    def _extract_plaintext_enhanced(self, file_path: str) -> Tuple[str, str]:
        """Extract plain text file with enhanced encoding detection."""
        encodings_to_try = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings_to_try:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                
                if text and len(text.strip()) > 10:
                    return text, f"txt:{encoding}"
                    
            except UnicodeDecodeError:
                continue
            except Exception as e:
                if self.verbose:
                    logger.warning(f"Failed to read with {encoding}: {e}")
                continue
        
        # Final fallback with error handling
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            return text, "txt:utf8_ignore"
        except Exception as e:
            logger.error(f"Failed to read plain text file: {e}")
            return "", "txt:error"
    
    def _extract_plaintext(self, file_path: str) -> Tuple[str, str]:
        """Extract plain text file (legacy method)."""
        try:
            # Try UTF-8 first
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            return text, "txt:utf8"
        except UnicodeDecodeError:
            # Try with error handling
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
                return text, "txt:utf8_ignore"
            except Exception as e:
                logger.error(f"Failed to read plain text file: {e}")
                return "", "txt:error"


# Convenience function
def extract_text_from_file_unified(file_path: str, verbose: bool = False) -> Tuple[str, str]:
    """
    Extract text from any supported file format with enhanced processing.
    
    Args:
        file_path: Path to file
        verbose: Enable verbose logging
        
    Returns:
        Tuple of (extracted_text, method_used)
    """
    extractor = UnifiedTextExtractor(verbose=verbose)
    return extractor.extract_text(file_path)


# Enhanced convenience function with caching
def extract_text_from_file_unified_enhanced(file_path: str, verbose: bool = False) -> Tuple[str, str]:
    """
    Extract text from any supported file format with all enhanced features.
    
    Args:
        file_path: Path to file
        verbose: Enable verbose logging
        
    Returns:
        Tuple of (extracted_text, method_used)
    """
    extractor = UnifiedTextExtractor(verbose=verbose)
    return extractor.extract_text(file_path)


# Backward compatibility wrapper
def extract_text_from_file_smart(file_path: str) -> str:
    """
    Backward compatibility wrapper.
    Returns only text (not method) for existing code.
    """
    text, method = extract_text_from_file_unified(file_path, verbose=False)
    return text
