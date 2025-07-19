"""
DEPRECATED: This module is deprecated and will be removed in a future version.

Use src.document_processing_unified instead for all document processing needs.

This module is kept only for backward compatibility with test files.
Production code should use document_processing_unified.
"""

import warnings
warnings.warn(
    "document_processing is deprecated. Use document_processing_unified instead.",
    DeprecationWarning,
    stacklevel=2
)

import os
import logging
import tempfile
from pathlib import Path
from typing import Dict, Any, Optional, Tuple, List
import requests
from urllib.parse import urlparse
import re
import time
import warnings

# Import text extraction libraries with better error handling
try:
    from bs4 import BeautifulSoup
    BEAUTIFULSOUP_AVAILABLE = True
except ImportError:
    BEAUTIFULSOUP_AVAILABLE = False
    logging.warning("BeautifulSoup not available - install with: pip install beautifulsoup4")

try:
    from striprtf.striprtf import rtf_to_text
    STRIPRTF_AVAILABLE = True
except ImportError:
    STRIPRTF_AVAILABLE = False
    logging.warning("striprtf not available - install with: pip install striprtf")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available - install with: pip install python-docx")

# Import enhanced processing modules
from src.document_processing_unified import extract_text_from_file
from src.case_name_extraction_core import extract_case_name_and_date_comprehensive
try:
    from .unified_citation_processor_v2 import UnifiedCitationProcessorV2, ProcessingConfig
    from .unified_citation_processor_v2 import UnifiedCitationProcessorV2 as UnifiedCitationProcessor
except ImportError:
    from unified_citation_processor_v2 import UnifiedCitationProcessorV2, ProcessingConfig
    from unified_citation_processor_v2 import UnifiedCitationProcessorV2 as UnifiedCitationProcessor
try:
    from .app_final_vue import verify_citation_with_extraction
except ImportError:
    from app_final_vue import verify_citation_with_extraction
try:
    from .document_processing_unified import extract_text_from_url as unified_extract_text_from_url
except ImportError:
    from document_processing_unified import extract_text_from_url as unified_extract_text_from_url
try:
    from .citation_extractor import normalize_washington_citations
except ImportError:
    from citation_extractor import normalize_washington_citations

# Configure logger first
logger = logging.getLogger(__name__)

# Use the new unified citation processor V2
try:
    class EnhancedProcessor(UnifiedCitationProcessorV2):
        def __init__(self):
            config = ProcessingConfig(
                use_eyecite=True,
                use_regex=True,
                extract_case_names=True,
                extract_dates=True,
                enable_clustering=True,
                enable_deduplication=True,
                debug_mode=False
            )
            super().__init__(config)
        
        def extract_case_name_and_year(self, text: str, citation: str):
            """Extract case name and year using the unified processor"""
            # Use the unified processor's regex extraction
            results = self.process_text(text)
            
            # Find the specific citation
            for result in results:
                if result.citation == citation:
                    return {
                        'case_name': result.case_name,
                        'year': result.date
                    }
            
            return {'case_name': '', 'year': ''}
    
    ENHANCED_PROCESSOR_AVAILABLE = True
    logger.info("Using UnifiedCitationProcessorV2")
except ImportError as e:
    logger.warning(f"UnifiedCitationProcessorV2 not available: {e}")
    # Fallback to original processor
    try:
        class EnhancedProcessor(UnifiedCitationProcessorV2):
            def __init__(self):
                config = ProcessingConfig(
                    use_eyecite=True,
                    use_regex=True,
                    extract_case_names=True,
                    extract_dates=True,
                    enable_clustering=True,
                    enable_deduplication=True,
                    debug_mode=False
                )
                super().__init__(config)
            
            def extract_case_name_and_year(self, text: str, citation: str):
                """Extract case name and year using the unified processor"""
                # Use the unified processor's regex extraction
                results = self.process_text(text)
                
                # Find the specific citation
                for result in results:
                    if result.citation == citation:
                        return {
                            'case_name': result.case_name,
                            'year': result.date
                        }
                
                return {'case_name': '', 'year': ''}
        
        ENHANCED_PROCESSOR_AVAILABLE = True
        logger.info("Using fallback UnifiedCitationProcessor")
    except ImportError as e2:
        logger.warning(f"Fallback processor not available: {e2}")
        ENHANCED_PROCESSOR_AVAILABLE = False

class EnhancedDocumentProcessor:
    """
    Enhanced document processor implementing research-based best practices.
    """
    
    def __init__(self):
        if ENHANCED_PROCESSOR_AVAILABLE:
            self.processor = EnhancedProcessor()
        else:
            # Import the fallback processor
            self.processor = UnifiedCitationProcessor()
            logger.info("Using fallback processor")
        
        # OCR error correction patterns based on research
        self.ocr_corrections = {
            'common_errors': [
                (r'(\d+)\s*F\s*(\d+)d\s*(\d+)', r'\1 F.\2d \3'),  # Fix missing periods
                (r'(\d+)\s*U\.S\.C\s*(\d+)', r'\1 U.S.C. § \2'),   # Fix missing section
                (r'(\d+)\s*F\s*Supp\s*(\d+)', r'\1 F. Supp. \2'),  # Fix abbreviations
                (r'(\d+)\s*Wn\s*(\d+)d\s*(\d+)', r'\1 Wn.\2d \3'), # Fix Washington citations
                (r'(\d+)\s*P\s*(\d+)d\s*(\d+)', r'\1 P.\2d \3'),   # Fix Pacific citations
            ],
            'character_corrections': {
                'l': '1',  # lowercase L confused with 1
                'O': '0',  # uppercase O confused with 0
                'rn': 'm', # rn combination confused with m
                'vv': 'w', # double v confused with w
                'cl': 'd'  # cl combination confused with d
            }
        }
    
    def preprocess_text(self, text: str, skip_ocr_correction: bool = False) -> str:
        """
        Preprocess text using research-based cleaning and OCR correction.
        """
        if not text:
            return ""
        
        # Step 1: Apply OCR corrections (skip for direct text submissions)
        if skip_ocr_correction:
            corrected_text = text
            logger.debug("Skipping OCR correction for direct text submission")
        else:
            corrected_text = self._apply_ocr_corrections(text)
        
        # Step 2: Use enhanced text cleaning
        if ENHANCED_PROCESSOR_AVAILABLE:
            # Use the processor's text cleaning method
            try:
                cleaned_text = self.processor._clean_text(corrected_text)
            except AttributeError:
                # Fallback if _clean_text method doesn't exist
                cleaned_text = re.sub(r'\s+', ' ', corrected_text).strip()
        else:
            # Fallback cleaning
            cleaned_text = re.sub(r'\s+', ' ', corrected_text).strip()
        
        logger.debug(f"Text preprocessing: {len(text)} -> {len(cleaned_text)} characters")
        return cleaned_text
    
    def _apply_ocr_corrections(self, text: str) -> str:
        """Apply OCR error corrections based on research findings."""
        corrected = text
        
        # Apply pattern-based corrections
        for pattern, replacement in self.ocr_corrections['common_errors']:
            corrected = re.sub(pattern, replacement, corrected)
        
        # Apply character-level corrections in numeric contexts
        for error, correction in self.ocr_corrections['character_corrections'].items():
            # Only apply in numeric/citation contexts
            corrected = re.sub(rf'\b{error}(?=\d)', correction, corrected)
        
        return corrected

    def extract_text_from_html_file(self, file_path: str) -> str:
        """Enhanced HTML text extraction with better error handling."""
        if not BEAUTIFULSOUP_AVAILABLE:
            raise ImportError("BeautifulSoup is required for HTML processing. Install with: pip install beautifulsoup4")
        
        encodings_to_try = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings_to_try:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    html_content = f.read()
                
                # Parse HTML with BeautifulSoup
                soup = BeautifulSoup(html_content, 'html.parser')
                
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                
                # Extract text with proper spacing
                text = soup.get_text(separator='\n', strip=True)
                
                # Clean up extra whitespace
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = '\n'.join(chunk for chunk in chunks if chunk)
                
                logger.info(f"Successfully extracted {len(text)} characters from HTML file using {encoding}")
                return text
                
            except UnicodeDecodeError:
                logger.debug(f"Failed to decode HTML file with {encoding}, trying next encoding")
                continue
            except Exception as e:
                logger.error(f"Error processing HTML file with {encoding}: {str(e)}")
                if encoding == encodings_to_try[-1]:  # Last encoding
                    raise
                continue
        
        raise Exception("Could not decode HTML file with any supported encoding")

    def extract_text_from_rtf_file(self, file_path: str) -> str:
        """Enhanced RTF text extraction."""
        if not STRIPRTF_AVAILABLE:
            raise ImportError("striprtf is required for RTF processing. Install with: pip install striprtf")
        
        encodings_to_try = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings_to_try:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    rtf_content = f.read()
                
                # Convert RTF to plain text
                text = rtf_to_text(rtf_content)
                text = text.strip()
                
                logger.info(f"Successfully extracted {len(text)} characters from RTF file using {encoding}")
                return text
                
            except UnicodeDecodeError:
                logger.debug(f"Failed to decode RTF file with {encoding}, trying next encoding")
                continue
            except Exception as e:
                logger.error(f"Error processing RTF file with {encoding}: {str(e)}")
                if encoding == encodings_to_try[-1]:
                    raise
                continue
        
        raise Exception("Could not decode RTF file with any supported encoding")

    def extract_text_from_docx_file(self, file_path: str) -> str:
        """Enhanced DOCX text extraction."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")
        
        try:
            doc = docx.Document(file_path)
            paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
            text = "\n".join(paragraphs)
            
            logger.info(f"Successfully extracted {len(text)} characters from DOCX file")
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX file: {str(e)}")
            raise

    def extract_text_from_file(self, file_path: str) -> str:
        """
        Enhanced file text extraction with comprehensive error handling.
        """
        file_extension = Path(file_path).suffix.lower()
        
        try:
            if file_extension == '.pdf':
                logger.info("Processing PDF file with enhanced handler")
                text = extract_text_from_file(file_path)
                if text.startswith("Error:"):
                    raise Exception(text)
                return text
                
            elif file_extension in ['.doc', '.docx']:
                logger.info("Processing Word document")
                return self.extract_text_from_docx_file(file_path)
                
            elif file_extension in ['.html', '.htm']:
                logger.info("Processing HTML/HTM file")
                return self.extract_text_from_html_file(file_path)
                
            elif file_extension == '.rtf':
                logger.info("Processing RTF file")
                return self.extract_text_from_rtf_file(file_path)
                
            elif file_extension == '.txt':
                logger.info("Processing text file")
                encodings_to_try = ['utf-8', 'latin-1', 'cp1252']
                for encoding in encodings_to_try:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            text = f.read()
                        logger.info(f"Successfully read text file with {encoding}")
                        return text
                    except UnicodeDecodeError:
                        if encoding == encodings_to_try[-1]:
                            raise
                        continue
                        
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
                
        except Exception as e:
            logger.error(f"Error extracting text from file {file_path}: {str(e)}")
            raise

    def extract_text_from_url(self, url: str) -> str:
        """
        Enhanced URL text extraction with better error handling and content type detection.
        """
        try:
            logger.info(f"Fetching content from URL: {url}")
            
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
                'Accept-Language': 'en-US,en;q=0.5',
                'Accept-Encoding': 'gzip, deflate',
                'Connection': 'keep-alive',
            }
            
            # Use session for better connection handling
            session = requests.Session()
            session.headers.update(headers)
            
            response = session.get(url, timeout=30, stream=True)
            response.raise_for_status()
            
            content_type = response.headers.get('content-type', '').lower()
            logger.info(f"Content type: {content_type}")
            
            # Enhanced content type mapping
            content_type_handlers = {
                'application/pdf': ('.pdf', self._handle_pdf_url),
                'application/msword': ('.doc', self._handle_doc_url),
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document': ('.docx', self._handle_doc_url),
                'application/rtf': ('.rtf', self._handle_rtf_url),
                'text/html': (None, self._handle_html_url),
                'text/plain': (None, self._handle_text_url),
            }
            
            # Find appropriate handler
            handler = None
            for ct_pattern, (ext, handler_func) in content_type_handlers.items():
                if ct_pattern in content_type:
                    handler = handler_func
                    file_ext = ext
                    break
            
            if handler:
                if file_ext:
                    # Handle as document - download to temp file
                    return self._download_and_process(response, file_ext, handler)
                else:
                    # Handle as text/html directly
                    return handler(response)
            else:
                # Default to text handling
                return self._handle_text_url(response)
                
        except requests.RequestException as e:
            logger.error(f"Network error fetching URL {url}: {str(e)}")
            raise
        except Exception as e:
            logger.error(f"Error extracting text from URL {url}: {str(e)}")
            raise

    def _download_and_process(self, response, file_ext, handler):
        """Download URL content to temp file and process."""
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as temp_file:
            for chunk in response.iter_content(chunk_size=8192):
                temp_file.write(chunk)
            temp_file_path = temp_file.name
        
        try:
            text = self.extract_text_from_file(temp_file_path)
            logger.info(f"Successfully extracted {len(text)} characters from downloaded file")
            return text
        finally:
            os.unlink(temp_file_path)

    def _handle_html_url(self, response):
        """Handle HTML URL content."""
        if BEAUTIFULSOUP_AVAILABLE:
            soup = BeautifulSoup(response.text, 'html.parser')
            for script in soup(["script", "style"]):
                script.decompose()
            text = soup.get_text(separator='\n', strip=True)
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
        else:
            text = response.text
        return text

    def _handle_text_url(self, response):
        """Handle plain text URL content."""
        return response.text

    def _handle_pdf_url(self, temp_file_path):
        """Handle PDF URL content."""
        return self.extract_text_from_file(temp_file_path)

    def _handle_doc_url(self, temp_file_path):
        """Handle DOC/DOCX URL content."""
        return self.extract_text_from_file(temp_file_path)

    def _handle_rtf_url(self, temp_file_path):
        """Handle RTF URL content."""
        return self.extract_text_from_file(temp_file_path)

    def process_document(self, 
                        content: str = None,
                        file_path: str = None,
                        url: str = None,
                        extract_case_names: bool = True,
                        debug_mode: bool = False) -> Dict[str, Any]:
        """
        Enhanced document processing with research-based improvements.
        """
        start_time = time.time()
        
        try:
            # Step 1: Extract text based on input type
            if content:
                logger.info("Processing direct text content")
                text = content
                source_type = "text"
                source_name = "pasted_text"
            elif file_path:
                logger.info(f"Processing file: {file_path}")
                text = self.extract_text_from_file(file_path)
                source_type = "file"
                source_name = Path(file_path).name
            elif url:
                logger.info(f"Processing URL: {url}")
                text = self.extract_text_from_url(url)
                source_type = "url"
                source_name = urlparse(url).netloc
            else:
                raise ValueError("Must provide content, file_path, or url")

            if not text.strip():
                logger.error("No text content extracted from input")
                return {
                    "success": False,
                    "error": "No text content extracted",
                    "citations": [],
                    "case_names": []
                }

            # Step 2: Preprocess text with research-based improvements
            # Skip OCR correction for direct text submissions to avoid corruption
            skip_ocr = source_type == 'text' and not file_path and not url
            preprocessed_text = self.preprocess_text(text, skip_ocr_correction=skip_ocr)
            logger.info(f"Text preprocessing: {len(text)} -> {len(preprocessed_text)} characters")

            # Step 3: Extract citations using enhanced processor
            if ENHANCED_PROCESSOR_AVAILABLE:
                logger.info("Using enhanced citation extraction with UnifiedCitationProcessorV2")
                
                # Disable OCR correction in the processor for direct text submissions
                if source_type == 'text':
                    try:
                        self.processor.enable_ocr_correction(False)
                        logger.debug("Disabled OCR correction in processor for direct text submission")
                    except AttributeError:
                        logger.debug("OCR correction control not available in this processor")
                
                processing_result = self.processor.process_text(preprocessed_text, extract_case_names=extract_case_names)
                
                # Re-enable OCR correction for future use
                if source_type == 'text':
                    try:
                        self.processor.enable_ocr_correction(True)
                        logger.debug("Re-enabled OCR correction in processor")
                    except AttributeError:
                        logger.debug("OCR correction control not available in this processor")
                
                citations = processing_result.get('results', [])
                statistics = processing_result.get('statistics')
                
                # Convert CitationResult objects to dictionaries for compatibility
                formatted_citations = []
                for citation in citations:
                    if hasattr(citation, '__dict__'):
                        # It's a CitationResult object
                        citation_dict = {
                            'citation': citation.citation,
                            'case_name': citation.case_name,
                            'extracted_case_name': citation.extracted_case_name,
                            'canonical_name': citation.canonical_name,
                            'extracted_date': citation.extracted_date,
                            'canonical_date': citation.canonical_date,
                            'verified': citation.verified,
                            'court': citation.court,
                            'confidence': citation.confidence,
                            'method': citation.method,
                            'context': citation.context,
                            'is_parallel': citation.is_parallel,
                            'parallel_citations': citation.parallel_citations,
                            'url': citation.url,
                            'source': citation.source
                        }
                    else:
                        # It's already a dictionary
                        citation_dict = citation
                    
                    formatted_citations.append(citation_dict)
                
                # Step 3.5: Apply enhanced verification with extraction
                logger.info("Applying enhanced verification with extraction")
                formatted_citations = self.verify_citations(formatted_citations, preprocessed_text)
                
            else:
                # Fallback to original extraction
                logger.warning("Using fallback citation extraction")
                formatted_citations = self._fallback_extraction(preprocessed_text)
                statistics = {'total_citations': len(formatted_citations)}
                
                # Apply enhanced verification with extraction
                logger.info("Applying enhanced verification with extraction")
                formatted_citations = self.verify_citations(formatted_citations, preprocessed_text)

            # Step 4: Extract case names
            case_names = []
            if extract_case_names:
                case_names = self._extract_case_names(formatted_citations, preprocessed_text)

            # Step 5: Debug mode
            debug_info = {}
            if debug_mode and formatted_citations:
                citation_texts = [c.get('citation', '') for c in formatted_citations]
                debug_info = debug_extraction_pipeline(preprocessed_text, citation_texts)

            processing_time = time.time() - start_time
            
            result = {
                "success": True,
                "text_length": len(text),
                "processed_text_length": len(preprocessed_text),
                "source_type": source_type,
                "source_name": source_name,
                "citations": formatted_citations,
                "case_names": case_names,
                "statistics": statistics,
                "processing_time": processing_time,
                "ocr_corrections_applied": len(text) != len(preprocessed_text),
                "debug_info": debug_info if debug_mode else None
            }
            
            logger.info(f"✅ Document processing completed in {processing_time:.2f}s")
            logger.info(f"   Found {len(formatted_citations)} citations, {len(case_names)} case names")
            
            return result

        except Exception as e:
            logger.error(f"Error processing document: {str(e)}", exc_info=True)
            return {
                "success": False,
                "error": str(e),
                "citations": [],
                "case_names": [],
                "processing_time": time.time() - start_time
            }

    def _fallback_extraction(self, text: str) -> List[Dict]:
        """Fallback extraction using original processor."""
        try:
            # DEPRECATED: from .citation_extractor import CitationExtractor
            extractor = CitationExtractor(
                use_eyecite=True,
                use_regex=True,
                context_window=1000,
                deduplicate=True,
                extract_case_names=True
            )
            return extractor.extract(text)
        except Exception as e:
            logger.error(f"Fallback extraction failed: {e}")
            return []

    def _extract_case_names(self, citations: List[Dict], text: str) -> List[str]:
        """Extract unique case names from citations and text."""
        case_names = set()
        
        # Extract from citations
        for citation in citations:
            for field in ['case_name', 'extracted_case_name', 'canonical_name']:
                name = citation.get(field)
                if name and name != "N/A" and len(name) > 3:
                    case_names.add(name)
        
        # If no case names found, try direct extraction
        if not case_names:
            try:
                # DEPRECATED: from .citation_extractor import CitationExtractor
                extractor = CitationExtractor()
                text_case_names = extractor._extract_case_names_from_text(text)
                case_names.update(text_case_names)
            except Exception as e:
                logger.warning(f"Direct case name extraction failed: {e}")
        
        return list(case_names)

    def verify_citations(self, citations: List[Dict], text: str) -> List[Dict]:
        """Enhanced verification with case name and date extraction using verify_citation_with_extraction."""
        from .app_final_vue import verify_citation_with_extraction
        
        for citation in citations:
            try:
                citation_text = citation.get('citation', '')
                if not citation_text:
                    continue
                
                # Use the verify_citation_with_extraction function for consistent processing
                verification_result = verify_citation_with_extraction(citation_text, text)
                
                # Update citation with extracted and canonical data
                citation.update({
                    'extracted_case_name': verification_result.get('extracted_case_name', 'N/A'),
                    'extracted_date': verification_result.get('extracted_date', 'N/A'),
                    'canonical_name': verification_result.get('canonical_name', 'N/A'),
                    'canonical_date': verification_result.get('canonical_date', 'N/A'),
                    'case_name': verification_result.get('extracted_case_name', 'N/A'),  # Use extracted as primary
                    'verified': verification_result.get('verified', False),
                    'source': verification_result.get('source', 'Unknown'),
                    'confidence': verification_result.get('confidence', 0.0),
                    'method': verification_result.get('method', 'enhanced_processor'),
                    'url': verification_result.get('url', ''),
                    'court': verification_result.get('court', 'N/A'),
                    'docket_number': verification_result.get('docket_number', 'N/A'),
                    'context': citation.get('context', text[:200] + '...' if len(text) > 200 else text)
                })
                
                # Store debug info
                citation['metadata'] = citation.get('metadata', {})
                citation['metadata'].update({
                    'verification_method': 'verify_citation_with_extraction',
                    'extraction_success': verification_result.get('extracted_case_name') != 'N/A',
                    'verification_success': verification_result.get('verified', False)
                })
                
                logger.debug(f"Enhanced verification for {citation_text}:")
                logger.debug(f"  Extracted: {citation['extracted_case_name']} ({citation['extracted_date']})")
                logger.debug(f"  Canonical: {citation['canonical_name']} ({citation['canonical_date']})")
                    
            except Exception as e:
                logger.warning(f"Enhanced extraction failed for {citation.get('citation', '')}: {e}")
                # Set default values on error
                citation.update({
                    'extracted_case_name': 'N/A',
                    'extracted_date': 'N/A',
                    'canonical_name': 'N/A',
                    'canonical_date': 'N/A',
                    'case_name': 'N/A',
                    'verified': False,
                    'source': 'error',
                    'confidence': 0.0,
                    'method': 'error',
                    'error': str(e)
                })
                
        return citations

# Create global instance
enhanced_processor = EnhancedDocumentProcessor()

# Convenience functions for backward compatibility
def process_document(content: str = None, file_path: str = None, url: str = None, 
                    extract_case_names: bool = True, debug_mode: bool = False) -> Dict[str, Any]:
    """Process document using enhanced processor."""
    return enhanced_processor.process_document(
        content=content, 
        file_path=file_path, 
        url=url, 
        extract_case_names=extract_case_names,
        debug_mode=debug_mode
    )

def extract_text_from_file(file_path: str) -> str:
    """
    DEPRECATED: Use src.document_processing_unified.extract_text_from_file instead.
    This function will be removed in a future version.
    """
    warnings.warn(
        "extract_text_from_file is deprecated. Use src.document_processing_unified.extract_text_from_file instead.",
        DeprecationWarning,
        stacklevel=2
    )
    from src.document_processing_unified import extract_text_from_file as unified_extract_text_from_file
    return unified_extract_text_from_file(file_path)

def extract_text_from_url(url: str) -> str:
    """
    DEPRECATED: Use src.document_processing_unified.extract_text_from_url instead.
    This function will be removed in a future version.
    """
    warnings.warn(
        "extract_text_from_url is deprecated. Use src.document_processing_unified.extract_text_from_url instead.",
        DeprecationWarning,
        stacklevel=2
    )
    from src.document_processing_unified import extract_text_from_url as unified_extract_text_from_url
    return unified_extract_text_from_url(url)

# For backward compatibility with existing verification function
def verify_citations_with_contamination_fix(citations: List[Dict], text: str) -> List[Dict]:
    import logging
    logger = logging.getLogger("verification")
    logger.info(f"[VERIFY] Starting contamination-free verification of {len(citations)} citations")
    verified_citations = []
    for citation_data in citations:
        citation_text = citation_data.get('citation') or citation_data.get('citation_text', '')
        if not citation_text:
            continue
        if citation_data.get('is_parallel_group', False):
            logger.info(f"[VERIFY] Skipping combined citation: {citation_text}")
            continue
        logger.debug(f"[VERIFY] Processing: {citation_text}")
        extraction_result = extract_case_name_and_date(text=text, citation=citation_text)
        clean_result = {
            'citation': citation_text,
            'verified': extraction_result.get('verified', False),
            'extracted_case_name': extraction_result.get('case_name', 'N/A'),
            'extracted_date': extraction_result.get('extracted_date', 'N/A'),
            'case_name': extraction_result.get('case_name', 'N/A'),
            'canonical_name': extraction_result.get('canonical_name', 'N/A'),
            'canonical_date': extraction_result.get('canonical_date', 'N/A'),
            'source': extraction_result.get('source', 'Unknown'),
            'confidence': extraction_result.get('case_name_confidence', 0.0),
            'method': extraction_result.get('case_name_method', 'none'),
            'url': extraction_result.get('url', ''),
            'court': extraction_result.get('court', ''),
            'docket_number': extraction_result.get('docket_number', ''),
            'context': citation_data.get('context', text[:200] + '...' if len(text) > 200 else text),
            'verified_by_parallel': False,
            'debug_info': extraction_result.get('debug_info', {}),
            'contamination_check': extraction_result.get('debug_info', {}).get('contamination_check', 'unknown')
        }
        logger.debug(f"[VERIFY] Results for {citation_text}:")
        logger.debug(f"  Extracted: {clean_result['extracted_case_name']} ({clean_result['extracted_date']})")
        logger.debug(f"  Canonical: {clean_result['canonical_name']} ({clean_result['canonical_date']})")
        logger.debug(f"  Contamination check: {clean_result['contamination_check']}")
        verified_citations.append(clean_result)
    parallel_groups = {}
    for citation in verified_citations:
        group_id = citation.get('parallel_group_id')
        if group_id is not None:
            if group_id not in parallel_groups:
                parallel_groups[group_id] = []
            parallel_groups[group_id].append(citation)
    for group_id, group_citations in parallel_groups.items():
        any_verified = any(citation.get('verified', False) for citation in group_citations)
        if any_verified:
            for citation in group_citations:
                if not citation.get('verified', False):
                    citation['verified_by_parallel'] = True
                    citation['verified'] = True
                    citation['source'] = 'verified_by_parallel'
    logger.info(f"[VERIFY] Completed contamination-free verification: {len(verified_citations)} citations processed")
    return verified_citations

# Update the main verification logic to use the contamination-free version
class EnhancedDocumentProcessor:
    """
    Enhanced document processor implementing research-based best practices.
    """
    
    def __init__(self):
        if ENHANCED_PROCESSOR_AVAILABLE:
            self.processor = EnhancedProcessor()
        else:
            # Import the fallback processor
            self.processor = UnifiedCitationProcessor()
            logger.info("Using fallback processor")
        
        # OCR error correction patterns based on research
        self.ocr_corrections = {
            'common_errors': [
                (r'(\d+)\s*F\s*(\d+)d\s*(\d+)', r'\1 F.\2d \3'),  # Fix missing periods
                (r'(\d+)\s*U\.S\.C\s*(\d+)', r'\1 U.S.C. § \2'),   # Fix missing section
                (r'(\d+)\s*F\s*Supp\s*(\d+)', r'\1 F. Supp. \2'),  # Fix abbreviations
                (r'(\d+)\s*Wn\s*(\d+)d\s*(\d+)', r'\1 Wn.\2d \3'), # Fix Washington citations
                (r'(\d+)\s*P\s*(\d+)d\s*(\d+)', r'\1 P.\2d \3'),   # Fix Pacific citations
            ],
            'character_corrections': {
                'l': '1',  # lowercase L confused with 1
                'O': '0',  # uppercase O confused with 0
                'rn': 'm', # rn combination confused with m
                'vv': 'w', # double v confused with w
                'cl': 'd'  # cl combination confused with d
            }
        }
    
    def preprocess_text(self, text: str, skip_ocr_correction: bool = False) -> str:
        """
        Preprocess text using research-based cleaning and OCR correction.
        """
        if not text:
            return ""
        
        # Step 1: Apply OCR corrections (skip for direct text submissions)
        if skip_ocr_correction:
            corrected_text = text
            logger.debug("Skipping OCR correction for direct text submission")
        else:
            corrected_text = self._apply_ocr_corrections(text)
        
        # Step 2: Use enhanced text cleaning
        if ENHANCED_PROCESSOR_AVAILABLE:
            # Use the processor's text cleaning method
            try:
                cleaned_text = self.processor._clean_text(corrected_text)
            except AttributeError:
                # Fallback if _clean_text method doesn't exist
                cleaned_text = re.sub(r'\s+', ' ', corrected_text).strip()
        else:
            # Fallback cleaning
            cleaned_text = re.sub(r'\s+', ' ', corrected_text).strip()
        
        logger.debug(f"Text preprocessing: {len(text)} -> {len(cleaned_text)} characters")
        return cleaned_text
    
    def _apply_ocr_corrections(self, text: str) -> str:
        """Apply OCR error corrections based on research findings."""
        corrected = text
        
        # Apply pattern-based corrections
        for pattern, replacement in self.ocr_corrections['common_errors']:
            corrected = re.sub(pattern, replacement, corrected)
        
        # Apply character-level corrections in numeric contexts
        for error, correction in self.ocr_corrections['character_corrections'].items():
            # Only apply in numeric/citation contexts
            corrected = re.sub(rf'\b{error}(?=\d)', correction, corrected)
        
        return corrected

    def extract_text_from_html_file(self, file_path: str) -> str:
        """Enhanced HTML text extraction with better error handling."""
        if not BEAUTIFULSOUP_AVAILABLE:
            raise ImportError("BeautifulSoup is required for HTML processing. Install with: pip install beautifulsoup4")
        
        encodings_to_try = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings_to_try:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    html_content = f.read()
                
                # Parse HTML with BeautifulSoup
                soup = BeautifulSoup(html_content, 'html.parser')
                
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                
                # Extract text with proper spacing
                text = soup.get_text(separator='\n', strip=True)
                
                # Clean up extra whitespace
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = '\n'.join(chunk for chunk in chunks if chunk)
                
                logger.info(f"Successfully extracted {len(text)} characters from HTML file using {encoding}")
                return text
                
            except UnicodeDecodeError:
                logger.debug(f"Failed to decode HTML file with {encoding}, trying next encoding")
                continue
            except Exception as e:
                logger.error(f"Error processing HTML file with {encoding}: {str(e)}")
                if encoding == encodings_to_try[-1]:  # Last encoding
                    raise
                continue
        
        raise Exception("Could not decode HTML file with any supported encoding")

    def extract_text_from_rtf_file(self, file_path: str) -> str:
        """Enhanced RTF text extraction."""
        if not STRIPRTF_AVAILABLE:
            raise ImportError("striprtf is required for RTF processing. Install with: pip install striprtf")
        
        encodings_to_try = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings_to_try:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    rtf_content = f.read()
                
                # Convert RTF to plain text
                text = rtf_to_text(rtf_content)
                text = text.strip()
                
                logger.info(f"Successfully extracted {len(text)} characters from RTF file using {encoding}")
                return text
                
            except UnicodeDecodeError:
                logger.debug(f"Failed to decode RTF file with {encoding}, trying next encoding")
                continue
            except Exception as e:
                logger.error(f"Error processing RTF file with {encoding}: {str(e)}")
                if encoding == encodings_to_try[-1]:
                    raise
                continue
        
        raise Exception("Could not decode RTF file with any supported encoding")

    def extract_text_from_docx_file(self, file_path: str) -> str:
        """Enhanced DOCX text extraction."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")
        
        try:
            doc = docx.Document(file_path)
            paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
            text = "\n".join(paragraphs)
            
            logger.info(f"Successfully extracted {len(text)} characters from DOCX file")
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX file: {str(e)}")
            raise

    def extract_text_from_file(self, file_path: str) -> str:
        """
        Enhanced file text extraction with comprehensive error handling.
        """
        file_extension = Path(file_path).suffix.lower()
        
        try:
            if file_extension == '.pdf':
                logger.info("Processing PDF file with enhanced handler")
                text = extract_text_from_file(file_path)
                if text.startswith("Error:"):
                    raise Exception(text)
                return text
                
            elif file_extension in ['.doc', '.docx']:
                logger.info("Processing Word document")
                return self.extract_text_from_docx_file(file_path)
                
            elif file_extension in ['.html', '.htm']:
                logger.info("Processing HTML/HTM file")
                return self.extract_text_from_html_file(file_path)
                
            elif file_extension == '.rtf':
                logger.info("Processing RTF file")
                return self.extract_text_from_rtf_file(file_path)
                
            elif file_extension == '.txt':
                logger.info("Processing text file")
                encodings_to_try = ['utf-8', 'latin-1', 'cp1252']
                for encoding in encodings_to_try:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            text = f.read()
                        logger.info(f"Successfully read text file with {encoding}")
                        return text
                    except UnicodeDecodeError:
                        if encoding == encodings_to_try[-1]:
                            raise
                        continue
                        
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
                
        except Exception as e:
            logger.error(f"Error extracting text from file {file_path}: {str(e)}")
            raise

    def extract_text_from_url(self, url: str) -> str:
        """
        Enhanced URL text extraction with better error handling and content type detection.
        """
        try:
            logger.info(f"Fetching content from URL: {url}")
            
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
                'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
                'Accept-Language': 'en-US,en;q=0.5',
                'Accept-Encoding': 'gzip, deflate',
                'Connection': 'keep-alive',
            }
            
            # Use session for better connection handling
            session = requests.Session()
            session.headers.update(headers)
            
            response = session.get(url, timeout=30, stream=True)
            response.raise_for_status()
            
            content_type = response.headers.get('content-type', '').lower()
            logger.info(f"Content type: {content_type}")
            
            # Enhanced content type mapping
            content_type_handlers = {
                'application/pdf': ('.pdf', self._handle_pdf_url),
                'application/msword': ('.doc', self._handle_doc_url),
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document': ('.docx', self._handle_doc_url),
                'application/rtf': ('.rtf', self._handle_rtf_url),
                'text/html': (None, self._handle_html_url),
                'text/plain': (None, self._handle_text_url),
            }
            
            # Find appropriate handler
            handler = None
            for ct_pattern, (ext, handler_func) in content_type_handlers.items():
                if ct_pattern in content_type:
                    handler = handler_func
                    file_ext = ext
                    break
            
            if handler:
                if file_ext:
                    # Handle as document - download to temp file
                    return self._download_and_process(response, file_ext, handler)
                else:
                    # Handle as text/html directly
                    return handler(response)
            else:
                # Default to text handling
                return self._handle_text_url(response)
                
        except requests.RequestException as e:
            logger.error(f"Network error fetching URL {url}: {str(e)}")
            raise
        except Exception as e:
            logger.error(f"Error extracting text from URL {url}: {str(e)}")
            raise

    def _download_and_process(self, response, file_ext, handler):
        """Download URL content to temp file and process."""
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as temp_file:
            for chunk in response.iter_content(chunk_size=8192):
                temp_file.write(chunk)
            temp_file_path = temp_file.name
        
        try:
            text = self.extract_text_from_file(temp_file_path)
            logger.info(f"Successfully extracted {len(text)} characters from downloaded file")
            return text
        finally:
            os.unlink(temp_file_path)

    def _handle_html_url(self, response):
        """Handle HTML URL content."""
        if BEAUTIFULSOUP_AVAILABLE:
            soup = BeautifulSoup(response.text, 'html.parser')
            for script in soup(["script", "style"]):
                script.decompose()
            text = soup.get_text(separator='\n', strip=True)
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)
        else:
            text = response.text
        return text

    def _handle_text_url(self, response):
        """Handle plain text URL content."""
        return response.text

    def _handle_pdf_url(self, temp_file_path):
        """Handle PDF URL content."""
        return self.extract_text_from_file(temp_file_path)

    def _handle_doc_url(self, temp_file_path):
        """Handle DOC/DOCX URL content."""
        return self.extract_text_from_file(temp_file_path)

    def _handle_rtf_url(self, temp_file_path):
        """Handle RTF URL content."""
        return self.extract_text_from_file(temp_file_path)

    def process_document(self, 
                        content: str = None,
                        file_path: str = None,
                        url: str = None,
                        extract_case_names: bool = True,
                        debug_mode: bool = False) -> Dict[str, Any]:
        """
        Enhanced document processing with research-based improvements.
        """
        start_time = time.time()
        
        try:
            # Step 1: Extract text based on input type
            if content:
                logger.info("Processing direct text content")
                text = content
                source_type = "text"
                source_name = "pasted_text"
            elif file_path:
                logger.info(f"Processing file: {file_path}")
                text = self.extract_text_from_file(file_path)
                source_type = "file"
                source_name = Path(file_path).name
            elif url:
                logger.info(f"Processing URL: {url}")
                text = self.extract_text_from_url(url)
                source_type = "url"
                source_name = urlparse(url).netloc
            else:
                raise ValueError("Must provide content, file_path, or url")

            if not text.strip():
                logger.error("No text content extracted from input")
                return {
                    "success": False,
                    "error": "No text content extracted",
                    "citations": [],
                    "case_names": []
                }

            # Step 2: Preprocess text with research-based improvements
            # Skip OCR correction for direct text submissions to avoid corruption
            skip_ocr = source_type == 'text' and not file_path and not url
            preprocessed_text = self.preprocess_text(text, skip_ocr_correction=skip_ocr)
            logger.info(f"Text preprocessing: {len(text)} -> {len(preprocessed_text)} characters")

            # Step 3: Extract citations using enhanced processor
            if ENHANCED_PROCESSOR_AVAILABLE:
                logger.info("Using enhanced citation extraction with UnifiedCitationProcessorV2")
                
                # Disable OCR correction in the processor for direct text submissions
                if source_type == 'text':
                    self.processor.enable_ocr_correction(False)
                    logger.debug("Disabled OCR correction in processor for direct text submission")
                
                processing_result = self.processor.process_text(preprocessed_text, extract_case_names=extract_case_names)
                
                # Re-enable OCR correction for future use
                if source_type == 'text':
                    self.processor.enable_ocr_correction(True)
                    logger.debug("Re-enabled OCR correction in processor")
                
                citations = processing_result.get('results', [])
                statistics = processing_result.get('statistics')
                
                # Convert CitationResult objects to dictionaries for compatibility
                formatted_citations = []
                for citation in citations:
                    if hasattr(citation, '__dict__'):
                        # It's a CitationResult object
                        citation_dict = {
                            'citation': citation.citation,
                            'case_name': citation.case_name,
                            'extracted_case_name': citation.extracted_case_name,
                            'canonical_name': citation.canonical_name,
                            'extracted_date': citation.extracted_date,
                            'canonical_date': citation.canonical_date,
                            'verified': citation.verified,
                            'court': citation.court,
                            'confidence': citation.confidence,
                            'method': citation.method,
                            'context': citation.context,
                            'is_parallel': citation.is_parallel,
                            'parallel_citations': citation.parallel_citations,
                            'url': citation.url,
                            'source': citation.source
                        }
                    else:
                        # It's already a dictionary
                        citation_dict = citation
                    
                    formatted_citations.append(citation_dict)
                
                # Step 3.5: Apply enhanced verification with extraction
                logger.info("Applying enhanced verification with extraction")
                formatted_citations = self.verify_citations(formatted_citations, preprocessed_text)
                
            else:
                # Fallback to original extraction
                logger.warning("Using fallback citation extraction")
                formatted_citations = self._fallback_extraction(preprocessed_text)
                statistics = {'total_citations': len(formatted_citations)}
                
                # Apply enhanced verification with extraction
                logger.info("Applying enhanced verification with extraction")
                formatted_citations = self.verify_citations(formatted_citations, preprocessed_text)

            # Step 4: Extract case names
            case_names = []
            if extract_case_names:
                case_names = self._extract_case_names(formatted_citations, preprocessed_text)

            # Step 5: Debug mode
            debug_info = {}
            if debug_mode and formatted_citations:
                citation_texts = [c.get('citation', '') for c in formatted_citations]
                debug_info = debug_extraction_pipeline(preprocessed_text, citation_texts)

            processing_time = time.time() - start_time
            
            result = {
                "success": True,
                "text_length": len(text),
                "processed_text_length": len(preprocessed_text),
                "source_type": source_type,
                "source_name": source_name,
                "citations": formatted_citations,
                "case_names": case_names,
                "statistics": statistics,
                "processing_time": processing_time,
                "ocr_corrections_applied": len(text) != len(preprocessed_text),
                "debug_info": debug_info if debug_mode else None
            }
            
            logger.info(f"✅ Document processing completed in {processing_time:.2f}s")
            logger.info(f"   Found {len(formatted_citations)} citations, {len(case_names)} case names")
            
            return result

        except Exception as e:
            logger.error(f"Error processing document: {str(e)}", exc_info=True)
            return {
                "success": False,
                "error": str(e),
                "citations": [],
                "case_names": [],
                "processing_time": time.time() - start_time
            }

    def _fallback_extraction(self, text: str) -> List[Dict]:
        """Fallback extraction using original processor."""
        try:
            # DEPRECATED: from .citation_extractor import CitationExtractor
            extractor = CitationExtractor(
                use_eyecite=True,
                use_regex=True,
                context_window=1000,
                deduplicate=True,
                extract_case_names=True
            )
            return extractor.extract(text)
        except Exception as e:
            logger.error(f"Fallback extraction failed: {e}")
            return []

    def _extract_case_names(self, citations: List[Dict], text: str) -> List[str]:
        """Extract unique case names from citations and text."""
        case_names = set()
        
        # Extract from citations
        for citation in citations:
            for field in ['case_name', 'extracted_case_name', 'canonical_name']:
                name = citation.get(field)
                if name and name != "N/A" and len(name) > 3:
                    case_names.add(name)
        
        # If no case names found, try direct extraction
        if not case_names:
            try:
                # DEPRECATED: from .citation_extractor import CitationExtractor
                extractor = CitationExtractor()
                text_case_names = extractor._extract_case_names_from_text(text)
                case_names.update(text_case_names)
            except Exception as e:
                logger.warning(f"Direct case name extraction failed: {e}")
        
        return list(case_names)

    def verify_citations(self, citations: List[Dict], text: str) -> List[Dict]:
        """Enhanced verification with case name and date extraction using verify_citation_with_extraction."""
        from .app_final_vue import verify_citation_with_extraction
        
        for citation in citations:
            try:
                citation_text = citation.get('citation', '')
                if not citation_text:
                    continue
                
                # Use the verify_citation_with_extraction function for consistent processing
                verification_result = verify_citation_with_extraction(citation_text, text)
                
                # Update citation with extracted and canonical data
                citation.update({
                    'extracted_case_name': verification_result.get('extracted_case_name', 'N/A'),
                    'extracted_date': verification_result.get('extracted_date', 'N/A'),
                    'canonical_name': verification_result.get('canonical_name', 'N/A'),
                    'canonical_date': verification_result.get('canonical_date', 'N/A'),
                    'case_name': verification_result.get('extracted_case_name', 'N/A'),  # Use extracted as primary
                    'verified': verification_result.get('verified', False),
                    'source': verification_result.get('source', 'Unknown'),
                    'confidence': verification_result.get('confidence', 0.0),
                    'method': verification_result.get('method', 'enhanced_processor'),
                    'url': verification_result.get('url', ''),
                    'court': verification_result.get('court', 'N/A'),
                    'docket_number': verification_result.get('docket_number', 'N/A'),
                    'context': citation.get('context', text[:200] + '...' if len(text) > 200 else text)
                })
                
                # Store debug info
                citation['metadata'] = citation.get('metadata', {})
                citation['metadata'].update({
                    'verification_method': 'verify_citation_with_extraction',
                    'extraction_success': verification_result.get('extracted_case_name') != 'N/A',
                    'verification_success': verification_result.get('verified', False)
                })
                
                logger.debug(f"Enhanced verification for {citation_text}:")
                logger.debug(f"  Extracted: {citation['extracted_case_name']} ({citation['extracted_date']})")
                logger.debug(f"  Canonical: {citation['canonical_name']} ({citation['canonical_date']})")
                    
            except Exception as e:
                logger.warning(f"Enhanced extraction failed for {citation.get('citation', '')}: {e}")
                # Set default values on error
                citation.update({
                    'extracted_case_name': 'N/A',
                    'extracted_date': 'N/A',
                    'canonical_name': 'N/A',
                    'canonical_date': 'N/A',
                    'case_name': 'N/A',
                    'verified': False,
                    'source': 'error',
                    'confidence': 0.0,
                    'method': 'error',
                    'error': str(e)
                })
                
        return citations

# Create global instance
enhanced_processor = EnhancedDocumentProcessor()

# Convenience functions for backward compatibility
def process_document(content: str = None, file_path: str = None, url: str = None, 
                    extract_case_names: bool = True, debug_mode: bool = False) -> Dict[str, Any]:
    """Process document using enhanced processor."""
    return enhanced_processor.process_document(
        content=content, 
        file_path=file_path, 
        url=url, 
        extract_case_names=extract_case_names,
        debug_mode=debug_mode
    )

def extract_text_from_file(file_path: str) -> str:
    """
    DEPRECATED: Use src.document_processing_unified.extract_text_from_file instead.
    This function will be removed in a future version.
    """
    warnings.warn(
        "extract_text_from_file is deprecated. Use src.document_processing_unified.extract_text_from_file instead.",
        DeprecationWarning,
        stacklevel=2
    )
    from src.document_processing_unified import extract_text_from_file as unified_extract_text_from_file
    return unified_extract_text_from_file(file_path)

def extract_text_from_url(url: str) -> str:
    """
    DEPRECATED: Use src.document_processing_unified.extract_text_from_url instead.
    This function will be removed in a future version.
    """
    warnings.warn(
        "extract_text_from_url is deprecated. Use src.document_processing_unified.extract_text_from_url instead.",
        DeprecationWarning,
        stacklevel=2
    )
    from src.document_processing_unified import extract_text_from_url as unified_extract_text_from_url
    return unified_extract_text_from_url(url)

# For backward compatibility with existing verification function
def verify_citations_with_fallback(citations: List[Dict], text: str) -> List[Dict]:
    """
    Maintain compatibility with existing verification logic.
    """
    from .unified_citation_processor_v2 import UnifiedCitationProcessorV2 as UnifiedCitationProcessor
    
    logger.info(f"[VERIFY] Starting verification of {len(citations)} citations")
    verifier = UnifiedCitationProcessor()
    
    verified_citations = []
    for citation_data in citations:
        citation_text = citation_data.get('citation', citation_data.get('citation_text', ''))
        if not citation_text:
            continue
            
        if citation_data.get('is_parallel_group', False):
            logger.info(f"[VERIFY] Skipping combined citation: {citation_text}")
            continue
        
        # Normalize citation
        normalized_citation = normalize_washington_citations(citation_text)
        
        # Extract additional information
        extracted_case_name = citation_data.get('case_name', '')
        extracted_date = citation_data.get('date', '')
        
        # Verify citation
        verification_result = verifier.verify_citation(
            normalized_citation,
            case_name=extracted_case_name,
            year=extracted_date
        )
        
        # Merge results - map UnifiedCitationProcessor fields to expected format
        merged_result = {
            **citation_data,
            'verified': verification_result.get('verified', False),
            'url': verification_result.get('sources', {}).get('courtlistener', {}).get('url', ''),
            'court': verification_result.get('sources', {}).get('courtlistener', {}).get('court', ''),
            'docket_number': verification_result.get('sources', {}).get('courtlistener', {}).get('docket_number', ''),
            'canonical_date': verification_result.get('sources', {}).get('courtlistener', {}).get('date_filed', ''),
            'source': verification_result.get('verified_by', 'Unknown'),
            'confidence': 0.9 if verification_result.get('verified', False) else 0.0,
            'error': verification_result.get('error', ''),
            'canonical_name': verification_result.get('sources', {}).get('courtlistener', {}).get('case_name', ''),
            'extracted_case_name': extracted_case_name,
            'extracted_date': extracted_date,
            'context': citation_data.get('context', text),
            'verified_by_parallel': False
        }
        verified_citations.append(merged_result)
    
    # Apply parallel citation fallback logic
    parallel_groups = {}
    for citation in verified_citations:
        group_id = citation.get('parallel_group_id')
        if group_id is not None:
            if group_id not in parallel_groups:
                parallel_groups[group_id] = []
            parallel_groups[group_id].append(citation)
    
    for group_id, group_citations in parallel_groups.items():
        any_verified = any(citation.get('verified', False) for citation in group_citations)
        if any_verified:
            for citation in group_citations:
                if not citation.get('verified', False):
                    citation['verified_by_parallel'] = True
                    citation['verified'] = True
                    citation['source'] = 'verified_by_parallel'
    
    # Add back combined citations
    for citation_data in citations:
        if citation_data.get('is_parallel_group', False):
            group_id = citation_data.get('parallel_group_id')
            if group_id is not None:
                component_verified = any(
                    citation.get('verified', False) 
                    for citation in verified_citations 
                    if citation.get('parallel_group_id') == group_id
                )
                
                combined_result = {
                    **citation_data,
                    'verified': component_verified,
                    'source': 'combined_display',
                    'verified_by_parallel': component_verified,
                    'context': citation_data.get('context', text)
                }
                verified_citations.append(combined_result)
    
    logger.info(f"[VERIFY] Completed verification: {len(verified_citations)} citations processed")
    return verified_citations

# Add the verify_citations_with_contamination_fix function as provided by the user

def extract_and_verify_citations(text: str, api_key: str = None) -> tuple:
    """
    Backward compatibility function for extract_and_verify_citations.
    Returns a tuple of (citations, metadata) for compatibility with existing tests.
    """
    processor = EnhancedDocumentProcessor()
    result = processor.process_document(content=text, extract_case_names=True)
    
    if result.get("success"):
        citations = result.get("citations", [])
        metadata = {
            "statistics": result.get("statistics", {}),
            "processing_time": result.get("processing_time", 0),
            "text_length": result.get("text_length", 0),
            "case_names": result.get("case_names", [])
        }
        return citations, metadata
    else:
        return [], {"error": result.get("error", "Unknown error")}
