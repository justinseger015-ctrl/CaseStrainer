# Import necessary modules
from flask import (
    Flask,
    send_from_directory,
    request,
    jsonify,
    render_template,
    redirect,
    url_for,
    session,
)
from config import (
    SECRET_KEY,
    DATABASE_FILE,
    UPLOAD_FOLDER,
    ALLOWED_EXTENSIONS,
    SESSION_TYPE,
    SESSION_FILE_DIR,
    SESSION_COOKIE_PATH,
    COURTLISTENER_API_KEY,
    configure_logging,
)
import os
import sys
import json
import sqlite3
import logging
import socket
import re
import requests
import traceback
import uuid
import tempfile
import datetime
import time
from flask_session import Session
from bs4 import BeautifulSoup
import urllib.parse
from werkzeug.utils import secure_filename
from werkzeug.middleware.proxy_fix import ProxyFix
from concurrent.futures import ThreadPoolExecutor
import threading
from flask_cors import CORS
from datetime import timedelta
from enhanced_validator_production import (
    enhanced_validator_bp as enhanced_validator_production_bp,
    register_enhanced_validator,
)
from citation_api import citation_api
from vue_api import api_blueprint
from vue_api_endpoints import vue_api

# Helper functions and constants remain at module level

# Global state to track processing progress
processing_state = {
    "total_citations": 0,
    "processed_citations": 0,
    "is_complete": False,
}

# Dictionary to store analysis results
analysis_results = {}

# Thread-local storage for API keys
thread_local = threading.local()

logger = logging.getLogger(__name__)


def create_app():
    configure_logging()
    app = Flask(
        __name__,
        static_folder=os.path.join(
            os.path.dirname(os.path.abspath(__file__)), "../casestrainer-vue/dist"
        ),
        template_folder=os.path.join(
            os.path.dirname(os.path.abspath(__file__)), "../casestrainer-vue/dist"
        ),
        static_url_path="",
    )

    # Register vue_api endpoints for enhanced validator and Vue.js API
    app.register_blueprint(vue_api, url_prefix="/casestrainer/api")

    # --- Global error handlers for consistent API responses ---
    @app.errorhandler(400)
    def bad_request(error):
        return (
            jsonify(
                {"status": "error", "message": "Bad request", "details": str(error)}
            ),
            400,
        )

    @app.errorhandler(404)
    def not_found(error):
        return (
            jsonify(
                {
                    "status": "error",
                    "message": "Resource not found",
                    "details": str(error),
                }
            ),
            404,
        )

    @app.errorhandler(500)
    def internal_error(error):
        return (
            jsonify(
                {
                    "status": "error",
                    "message": "Internal server error",
                    "details": str(error),
                }
            ),
            500,
        )

    # --- Session and Upload Config (from app.py) ---
    app.config["SECRET_KEY"] = SECRET_KEY or "devkey"
    app.config["SESSION_TYPE"] = "filesystem"
    app.config["PERMANENT_SESSION_LIFETIME"] = 3600 * 24
    app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
    app.config["SESSION_FILE_DIR"] = os.path.join(
        os.path.dirname(__file__), "..", "casestrainer_sessions"
    )
    app.config["SESSION_COOKIE_PATH"] = "/casestrainer/"

    # --- Optionally register vue_api blueprint ---
    try:
        from vue_api import api_blueprint
    except ImportError:
        api_blueprint = None
    if api_blueprint:
        app.register_blueprint(api_blueprint, url_prefix="/casestrainer/api")
        os.makedirs("logs", exist_ok=True)
        try:
            from concurrent_log_handler import ConcurrentRotatingFileHandler

            log_handler = ConcurrentRotatingFileHandler(
                "logs/casestrainer.log", maxBytes=5 * 1024 * 1024, backupCount=5
            )
        except ImportError:
            log_handler = logging.FileHandler("logs/casestrainer.log")

        logging.basicConfig(
            level=logging.DEBUG,  # Changed to DEBUG level for more detailed logs
            format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
            handlers=[log_handler, logging.StreamHandler()],
        )

    # Configure CORS with more specific settings
    CORS(
        app,
        resources={
            r"/*": {
                "origins": [
                    "https://wolf.law.uw.edu",
                    "http://localhost:5000",
                    "http://127.0.0.1:5000",
                ],
                "methods": ["GET", "POST", "PUT", "DELETE", "OPTIONS", "PATCH"],
                "allow_headers": [
                    "Content-Type",
                    "Authorization",
                    "X-Requested-With",
                    "X-CSRFToken",
                ],
                "expose_headers": ["Content-Type", "Content-Length", "X-CSRFToken"],
                "supports_credentials": True,
                "max_age": 600,  # Cache preflight requests for 10 minutes
            },
            "/analyze": {
                "origins": [
                    "https://wolf.law.uw.edu",
                    "http://localhost:5000",
                    "http://127.0.0.1:5000",
                ],
                "methods": ["POST", "OPTIONS"],
                "allow_headers": ["Content-Type", "Authorization"],
                "supports_credentials": True,
            },
        },
    )
    logger.info("CORS configured with specific origins for security")

    # Load SECRET_KEY with better error handling and security
    secret_key = os.environ.get("SECRET_KEY")
    if not secret_key:
        try:
            # Try to load from config.json in project root
            config_path = os.path.join(
                os.path.dirname(os.path.abspath(__file__)), "..", "config.json"
            )
            if os.path.exists(config_path):
                with open(config_path, "r") as f:
                    config = json.load(f)
                    secret_key = config.get("SECRET_KEY", "")
                    logger.info(
                        f"Loaded SECRET_KEY from config.json: {secret_key[:5]}..."
                        if secret_key
                        else "No SECRET_KEY found in config.json"
                    )
        except (FileNotFoundError, json.JSONDecodeError) as e:
            logger.error(f"Error loading config.json: {e}")

    # If still no secret key, generate a random one (not ideal for production but better than hardcoded)
    if not secret_key:
        import secrets

        secret_key = secrets.token_hex(32)
        logger.warning("Generated random SECRET_KEY - this will change on restart!")

    app.config["SECRET_KEY"] = secret_key

    # Session configuration with better security settings
    app.config["SESSION_TYPE"] = "filesystem"
    app.config["SESSION_FILE_DIR"] = os.path.join(
        os.path.dirname(os.path.abspath(__file__)), "flask_session"
    )
    os.makedirs(app.config["SESSION_FILE_DIR"], exist_ok=True)
    app.config["SESSION_COOKIE_SECURE"] = True
    app.config["SESSION_COOKIE_HTTPONLY"] = True
    app.config["SESSION_COOKIE_SAMESITE"] = "Lax"
    app.config["PERMANENT_SESSION_LIFETIME"] = timedelta(days=7)
    app.config["SESSION_USE_SIGNER"] = True  # Sign the session cookie

    # Initialize Flask-Session
    Session(app)
    logger.info("Flask-Session initialized with secure settings")

    # URL prefix configuration for Nginx proxy
    app_root = os.environ.get("APPLICATION_ROOT", "/casestrainer")
    app.config["APPLICATION_ROOT"] = app_root
    logger.info(f"Using APPLICATION_ROOT: {app_root}")

    # Configure for reverse proxy with proper settings for Nginx
    app.wsgi_app = ProxyFix(
        app.wsgi_app,
        x_proto=1,  # Number of proxy servers setting X-Forwarded-Proto
        x_host=1,  # Number of proxy servers setting X-Forwarded-Host
        x_port=1,  # Number of proxy servers setting X-Forwarded-Port
        x_prefix=1,  # Number of proxy servers setting X-Forwarded-Prefix
        x_for=1,  # Number of proxy servers setting X-Forwarded-For
    )
    logger.info("ProxyFix middleware configured for Nginx proxy")

    # Create thread pool for handling concurrent requests
    thread_pool = ThreadPoolExecutor(max_workers=10)

    # Import eyecite for citation extraction
    try:
        from eyecite import get_citations
        from eyecite.tokenizers import AhocorasickTokenizer

        try:
            tokenizer = AhocorasickTokenizer()
            EYECITE_AVAILABLE = True
            logger.info(
                "Eyecite library and AhocorasickTokenizer loaded successfully for citation extraction"
            )
        except ImportError as e:
            EYECITE_AVAILABLE = False
            tokenizer = None
            logger.warning(
                f"Eyecite not installed: {str(e)}. Using regex patterns for citation extraction."
            )
    except ImportError as e:
        EYECITE_AVAILABLE = False
        tokenizer = None
        logger.warning(
            f"Eyecite not installed: {str(e)}. Using regex patterns for citation extraction."
        )

    # Register the blueprints
    # Note: citation_api already includes the /api prefix in its route definitions
    app.register_blueprint(
        enhanced_validator_production_bp, url_prefix="/casestrainer/enhanced-validator"
    )
    app.register_blueprint(citation_api, url_prefix="/casestrainer/api")
    try:
        register_enhanced_validator(app)
    except Exception as e:
        print(f"Enhanced validator registration failed: {e}")

    # Print all registered Flask routes at startup for debugging
    try:
        from enhanced_validator_production import print_registered_routes

        print_registered_routes(app)
    except Exception as e:
        print(f"Could not print registered routes: {e}")

    # --- Error Handler for 500 ---
    @app.errorhandler(500)
    def server_error(error):
        return jsonify({"error": "Internal server error"}), 500

    # Direct route for analyze endpoint to handle file uploads
    @app.route("/api/analyze", methods=["POST", "OPTIONS"])
    def direct_analyze():
        # Log the request for debugging
        logger.info(
            f"[DEBUG] /api/analyze endpoint called with method: {request.method}"
        )
        logger.info(f"[DEBUG] Request headers: {dict(request.headers)}")
        logger.info(f"[DEBUG] Request files: {request.files}")
        logger.info(f"[DEBUG] Request form: {request.form}")
        logger.info(f"[DEBUG] Request data: {request.data}")

        # Forward to the Blueprint's analyze function
        from citation_api import analyze as citation_analyze

        return citation_analyze()

        # API endpoints for Vue.js frontend

        # Removed the direct @app.route('/api/upload') definition and function
        # All upload requests are now handled via the Blueprint at /casestrainer/api/upload

        # ... (rest of the code remains the same)
        logger.warning(f"Request files: {request.files}")
        logger.warning(f"Request form: {request.form}")
        return jsonify({"error": "No file part"}), 400

        file = request.files["file"]
        logger.info(
            f"File received: {file.filename}, Content-Type: {file.content_type}, Size: {file.content_length} bytes"
        )

        if file.filename == "":
            logger.warning("No file selected")
            return jsonify({"error": "No file selected"}), 400

        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filepath = os.path.join(UPLOAD_FOLDER, filename)
            file.save(filepath)

            file_size = os.path.getsize(filepath)
            logger.info(f"File saved to {filepath}, Size: {file_size} bytes")

            try:
                # Log start of citation extraction
                logger.info(f"Starting citation extraction from file: {filename}")
                start_time = time.time()

                # Extract citations from the file
                logger.info(f"Calling extract_citations_from_file for {filepath}")
                citations = extract_citations_from_file(filepath)
                logger.info(f"[DEBUG] Citations extracted: {len(citations)}")

                # Log extraction completion
                extraction_time = time.time() - start_time
                logger.info(
                    f"Citation extraction complete. Found {len(citations)} citations in {extraction_time:.2f} seconds"
                )

                # Log each citation found
                for i, citation in enumerate(citations):
                    logger.info(f"Citation {i+1}: {citation['text']}")

                # Process citations
                logger.info("Starting citation verification")
                verification_start = time.time()

                # Enable more verbose debugging
                print("\n==== STARTING CITATION VERIFICATION PROCESS ====\n")
                logger.info("==== STARTING CITATION VERIFICATION PROCESS ====")

                for i, citation in enumerate(citations):
                    print(
                        f"\n==== VERIFYING CITATION {i+1}/{len(citations)}: {citation['text']} ====\n"
                    )
                    logger.info(
                        f"==== VERIFYING CITATION {i+1}/{len(citations)}: {citation['text']} ===="
                    )

                    try:
                        # Get the first context if available
                        context = (
                            citation.get("contexts", [])[0].get("text", "")
                            if citation.get("contexts")
                            else ""
                        )
                        verification_result = verify_citation(
                            citation["text"], context=context
                        )

                        print(
                            f"\n==== VERIFICATION RESULT FOR CITATION {i+1}: {verification_result} ====\n"
                        )
                        logger.info(
                            f"==== VERIFICATION RESULT FOR CITATION {i+1}: {verification_result} ===="
                        )

                        # Update citation with verification results
                        citation["valid"] = verification_result.get("found", False)
                        citation["explanation"] = verification_result.get("explanation")

                        # Update metadata with source information
                        if "source" in verification_result:
                            citation["metadata"]["source"] = verification_result[
                                "source"
                            ]

                        # Add case name if available
                        if (
                            verification_result.get("case_name")
                            and not citation["name"]
                        ):
                            citation["name"] = verification_result["case_name"]

                        # Add any additional details
                        if verification_result.get("details"):
                            for key, value in verification_result["details"].items():
                                if value and key not in citation["metadata"]:
                                    citation["metadata"][key] = value

                        # Add URL if available
                        if verification_result.get("url"):
                            citation["metadata"]["url"] = verification_result["url"]

                        # Add is_westlaw flag if available
                        if "is_westlaw" in verification_result:
                            citation["metadata"]["is_westlaw"] = verification_result[
                                "is_westlaw"
                            ]

                        print(f"\n==== UPDATED CITATION {i+1}: {citation} ====\n")
                        logger.info(f"==== UPDATED CITATION {i+1}: {citation} ====")

                        # Add more detailed logging
                        print(f"Citation {i+1} validity: {citation['valid']}")
                        print(
                            f"Citation {i+1} source: {citation['metadata'].get('source', 'None')}"
                        )
                        print(
                            f"Citation {i+1} explanation: {citation.get('explanation', 'None')}"
                        )

                        logger.info(f"Citation {i+1} validity: {citation['valid']}")
                        logger.info(
                            f"Citation {i+1} source: {citation['metadata'].get('source', 'None')}"
                        )
                        logger.info(
                            f"Citation {i+1} explanation: {citation.get('explanation', 'None')}"
                        )
                    except Exception as e:
                        print(f"\n==== ERROR VERIFYING CITATION {i+1}: {str(e)} ====\n")
                        logger.error(
                            f"==== ERROR VERIFYING CITATION {i+1}: {str(e)} ===="
                        )
                        logger.error(traceback.format_exc())

                        # Set citation as invalid with error explanation
                        citation["valid"] = False
                        citation["explanation"] = f"Error during verification: {str(e)}"

                print("\n==== CITATION VERIFICATION PROCESS COMPLETE ====\n")
                logger.info("==== CITATION VERIFICATION PROCESS COMPLETE ====")

                verification_time = time.time() - verification_start
                logger.info(
                    f"Citation verification complete in {verification_time:.2f} seconds"
                )

                total_time = time.time() - start_time
                logger.info(f"Total processing time: {total_time:.2f} seconds")

                # Add debug logging for backend→frontend data structure and verification category counts
                verified_by_cl = [
                    c
                    for c in citations
                    if c.get("metadata", {}).get("source") == "CourtListener"
                ]
                verified_by_other = [
                    c
                    for c in citations
                    if c.get("metadata", {}).get("source")
                    and c.get("metadata", {}).get("source") != "CourtListener"
                ]
                unverified = [
                    c for c in citations if not c.get("metadata", {}).get("source")
                ]
                debug_summary = {
                    "totalCitations": len(citations),
                    "verifiedByCL": len(verified_by_cl),
                    "verifiedByOther": len(verified_by_other),
                    "unverified": len(unverified),
                }
                print("Data being sent to frontend:", debug_summary)
                logger.info(f"Data being sent to frontend: {debug_summary}")
                return jsonify(
                    {
                        "message": f"Successfully analyzed {len(citations)} citations in {filename}",
                        "citations": citations,
                        "processing_time": total_time,
                        "verification_summary": debug_summary,
                    }
                )
            except Exception as e:
                logger.error(f"Error processing file: {str(e)}")
                logger.error(traceback.format_exc())
                return (
                    jsonify(
                        {
                            "error": f"Error processing file: {str(e)}",
                            "traceback": traceback.format_exc(),
                        }
                    ),
                    500,
                )
        else:
            logger.warning(f"Invalid file type: {file.filename}")
            return jsonify({"error": "Invalid file type"}), 400

    @app.route("/api/text", methods=["POST"])
    def analyze_text():
        logger.info(f"[DEBUG] /api/text endpoint called from {request.remote_addr}")
        data = request.get_json()
        logger.info(f"[DEBUG] Incoming data: {data}")
        if not data or "text" not in data:
            logger.warning("No text provided in request")
            return jsonify({"error": "No text provided"}), 400

        text = data["text"]
        logger.info(f"[DEBUG] Text sample: {text[:300]}...")
        check_multiple_sources = data.get("checkMultipleSources", False)

        try:
            # Extract citations from text using both eyecite and regex, deduplicated
            citations = extract_all_citations(text, logger=logger)
            logger.info(
                f"Extracted {len(citations)} (deduplicated) citations from text."
            )

            # Process citations
            logger.info("Starting citation verification for text input")
            verification_start = time.time()

            # Enable more verbose debugging
            print("\n==== STARTING CITATION VERIFICATION PROCESS FOR TEXT INPUT ====\n")
            logger.info(
                "==== STARTING CITATION VERIFICATION PROCESS FOR TEXT INPUT ===="
            )

            for i, citation in enumerate(citations):
                print(
                    f"\n==== VERIFYING CITATION {i+1}/{len(citations)}: {citation['text']} ====\n"
                )
                logger.info(
                    f"==== VERIFYING CITATION {i+1}/{len(citations)}: {citation['text']} ===="
                )

                try:
                    # Get the first context if available
                    context = (
                        citation.get("contexts", [])[0].get("text", "")
                        if citation.get("contexts")
                        else ""
                    )
                    verification_result = verify_citation(
                        citation["text"], context=context
                    )

                    print(
                        f"\n==== VERIFICATION RESULT FOR CITATION {i+1}: {verification_result} ====\n"
                    )
                    logger.info(
                        f"==== VERIFICATION RESULT FOR CITATION {i+1}: {verification_result} ===="
                    )

                    # Update citation with verification results
                    citation["valid"] = verification_result.get("found", False)
                    citation["explanation"] = verification_result.get("explanation")

                    # Update metadata with source information
                    if "source" in verification_result:
                        citation["metadata"]["source"] = verification_result["source"]

                    # Add case name if available
                    if verification_result.get("case_name") and not citation["name"]:
                        citation["name"] = verification_result["case_name"]

                    # Add any additional details
                    if verification_result.get("details"):
                        for key, value in verification_result["details"].items():
                            if value and key not in citation["metadata"]:
                                citation["metadata"][key] = value

                    # Add URL if available
                    if verification_result.get("url"):
                        citation["metadata"]["url"] = verification_result["url"]

                    # Add is_westlaw flag if available
                    if "is_westlaw" in verification_result:
                        citation["metadata"]["is_westlaw"] = verification_result[
                            "is_westlaw"
                        ]

                    print(f"\n==== UPDATED CITATION {i+1}: {citation} ====\n")
                    logger.info(f"==== UPDATED CITATION {i+1}: {citation} ====")

                except Exception as e:
                    print(f"\n==== ERROR VERIFYING CITATION {i+1}: {str(e)} ====\n")
                    logger.error(f"==== ERROR VERIFYING CITATION {i+1}: {str(e)} ====")
                    logger.error(traceback.format_exc())

                    # Set citation as invalid with error explanation
                    citation["valid"] = False
                    citation["explanation"] = f"Error during verification: {str(e)}"

            print("\n==== CITATION VERIFICATION PROCESS COMPLETE ====\n")
            logger.info("==== CITATION VERIFICATION PROCESS COMPLETE ====")

            verification_time = time.time() - verification_start
            logger.info(
                f"Citation verification complete in {verification_time:.2f} seconds"
            )

            return jsonify(
                {
                    "message": f"Successfully analyzed {len(citations)} citations in the provided text",
                    "citations": citations,
                    "processing_time": verification_time,
                }
            )
        except Exception as e:
            logger.error(f"Error analyzing text: {str(e)}")
            logger.error(traceback.format_exc())
            return jsonify({"error": f"Error analyzing text: {str(e)}"}), 500

    @app.route("/api/url", methods=["POST", "OPTIONS"])
    def analyze_url():
        logger.info(
            f"[DEBUG] /api/url endpoint called. Remote addr: {request.remote_addr}"
        )
        try:
            data = request.get_json()
            logger.info(f"[DEBUG] Incoming request data: {data}")
        except Exception as e:
            logger.error(f"[DEBUG] Error parsing JSON in /api/url: {str(e)}")
            logger.error(traceback.format_exc())
            return jsonify({"error": "Malformed JSON", "details": str(e)}), 400
        if not data or "url" not in data:
            logger.warning("No URL provided in request")
            return jsonify({"error": "No URL provided"}), 400

        url = data["url"]
        logger.info(f"[DEBUG] URL to analyze: {url}")
        check_multiple_sources = data.get("checkMultipleSources", False)

        try:
            # Fetch content from URL
            response = requests.get(url, timeout=30)
            response.raise_for_status()

            # Extract text from HTML if it's an HTML page
            content_type = response.headers.get("Content-Type", "").lower()
            if "text/html" in content_type:
                soup = BeautifulSoup(response.text, "html.parser")
                text = soup.get_text()
            else:
                text = response.text

            # Extract citations from the text
            citations = extract_all_citations(text, logger=logger)

            # Process citations
            logger.info("Starting citation verification for URL input")
            verification_start = time.time()

            # Enable more verbose debugging
            print("\n==== STARTING CITATION VERIFICATION PROCESS FOR URL INPUT ====\n")
            logger.info(
                "==== STARTING CITATION VERIFICATION PROCESS FOR URL INPUT ===="
            )

            for i, citation in enumerate(citations):
                print(
                    f"\n==== VERIFYING CITATION {i+1}/{len(citations)}: {citation['text']} ====\n"
                )
                logger.info(
                    f"==== VERIFYING CITATION {i+1}/{len(citations)}: {citation['text']} ===="
                )

                try:
                    # Get the first context if available
                    context = (
                        citation.get("contexts", [])[0].get("text", "")
                        if citation.get("contexts")
                        else ""
                    )
                    verification_result = verify_citation(
                        citation["text"], context=context
                    )

                    print(
                        f"\n==== VERIFICATION RESULT FOR CITATION {i+1}: {verification_result} ====\n"
                    )
                    logger.info(
                        f"==== VERIFICATION RESULT FOR CITATION {i+1}: {verification_result} ===="
                    )

                    # Update citation with verification results
                    citation["valid"] = verification_result.get("found", False)
                    citation["explanation"] = verification_result.get("explanation")

                    # Update metadata with source information
                    if "source" in verification_result:
                        citation["metadata"]["source"] = verification_result["source"]

                    # Add case name if available
                    if verification_result.get("case_name") and not citation["name"]:
                        citation["name"] = verification_result["case_name"]

                    # Add any additional details
                    if verification_result.get("details"):
                        for key, value in verification_result["details"].items():
                            if value and key not in citation["metadata"]:
                                citation["metadata"][key] = value

                    # Add URL if available
                    if verification_result.get("url"):
                        citation["metadata"]["url"] = verification_result["url"]

                    # Add is_westlaw flag if available
                    if "is_westlaw" in verification_result:
                        citation["metadata"]["is_westlaw"] = verification_result[
                            "is_westlaw"
                        ]

                    print(f"\n==== UPDATED CITATION {i+1}: {citation} ====\n")
                    logger.info(f"==== UPDATED CITATION {i+1}: {citation} ====")

                except Exception as e:
                    print(f"\n==== ERROR VERIFYING CITATION {i+1}: {str(e)} ====\n")
                    logger.error(f"==== ERROR VERIFYING CITATION {i+1}: {str(e)} ====")
                    logger.error(traceback.format_exc())

                    # Set citation as invalid with error explanation
                    citation["valid"] = False
                    citation["explanation"] = f"Error during verification: {str(e)}"

            print("\n==== CITATION VERIFICATION PROCESS COMPLETE ====\n")
            logger.info("==== CITATION VERIFICATION PROCESS COMPLETE ====")

            verification_time = time.time() - verification_start
            logger.info(
                f"Citation verification complete in {verification_time:.2f} seconds"
            )

            return jsonify(
                {
                    "message": f"Successfully analyzed {len(citations)} citations from {url}",
                    "citations": citations,
                    "processing_time": verification_time,
                }
            )
        except requests.RequestException as e:
            logger.error(f"Error fetching URL {url}: {str(e)}")
            return jsonify({"error": f"Error fetching URL: {str(e)}"}), 500
        except Exception as e:
            logger.error(f"Error analyzing URL content: {str(e)}")
            logger.error(traceback.format_exc())
            return jsonify({"error": f"Error analyzing URL content: {str(e)}"}), 500

    # Configure logging for the app
    app.logger.handlers = logger.handlers
    app.logger.setLevel(logger.level)

    # The following custom static asset routes are now disabled, as all static assets (css, js, img, fonts) should be served directly from the Vue build output (dist) via Flask's static file handling.
    # @app.route('/css/<path:path>')
    # @app.route('/casestrainer/css/<path:path>')
    # def serve_vue_css(path):
    #     vue_dist_dir = os.path.join(os.path.dirname(__file__), 'static', 'vue', 'css')
    #     response = send_from_directory(vue_dist_dir, path)
    #     # Enhanced cache control to ensure the latest version is always displayed
    #     response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0, post-check=0, pre-check=0'
    #     response.headers['Pragma'] = 'no-cache'
    #     response.headers['Expires'] = '-1'
    #     # Add a timestamp to force refresh
    #     response.headers['Last-Modified'] = f"{datetime.datetime.now().strftime('%a, %d %b %Y %H:%M:%S GMT')}"
    #     return response

    # @app.route('/img/<path:path>')
    # @app.route('/casestrainer/img/<path:path>')
    # def serve_vue_img(path):
    #     vue_dist_dir = os.path.join(os.path.dirname(__file__), 'static', 'vue', 'img')
    #     response = send_from_directory(vue_dist_dir, path)
    #     # Enhanced cache control to ensure the latest version is always displayed
    #     response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0, post-check=0, pre-check=0'
    #     response.headers['Pragma'] = 'no-cache'
    #     response.headers['Expires'] = '-1'
    #     # Add a timestamp to force refresh
    #     response.headers['Last-Modified'] = f"{datetime.datetime.now().strftime('%a, %d %b %Y %H:%M:%S GMT')}"
    #     return response

    # @app.route('/fonts/<path:path>')
    # @app.route('/casestrainer/fonts/<path:path>')
    # def serve_vue_fonts(path):
    #     vue_dist_dir = os.path.join(os.path.dirname(__file__), 'static', 'vue', 'fonts')
    #     response = send_from_directory(vue_dist_dir, path)
    #     # Enhanced cache control to ensure the latest version is always displayed
    #     response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0, post-check=0, pre-check=0'
    #     response.headers['Pragma'] = 'no-cache'
    #     response.headers['Expires'] = '-1'
    #     # Add a timestamp to force refresh
    #     response.headers['Last-Modified'] = f"{datetime.datetime.now().strftime('%a, %d %b %Y %H:%M:%S GMT')}"
    #     return response

    @app.route("/original")
    @app.route("/casestrainer/original")
    def redirect_to_api():
        return redirect("/api/")

    @app.route("/original/")
    @app.route("/casestrainer/original/")
    def redirect_to_api_slash():
        return redirect("/api/")

    @app.route("/vue-enhanced-validator")
    @app.route("/casestrainer/vue-enhanced-validator")
    def redirect_from_vue_to_enhanced_validator():
        return redirect("/")

    @app.route("/enhanced-validator-link")
    @app.route("/casestrainer/enhanced-validator-link")
    def serve_enhanced_validator_link():
        return send_from_directory("static", "enhanced_validator_link.html")

    @app.route("/citation_verification_results.json")
    @app.route("/casestrainer/citation_verification_results.json")
    def serve_citation_verification_results():
        return send_from_directory(
            os.path.join(os.path.dirname(__file__), "static"),
            "citation_verification_results.json",
        )

    @app.route("/database_verification_results.json")
    @app.route("/casestrainer/database_verification_results.json")
    def serve_database_verification_results():
        return send_from_directory(
            os.path.join(os.path.dirname(__file__), "static"),
            "database_verification_results.json",
        )

    @app.route("/api/confirmed_with_multitool_data")
    @app.route("/api/confirmed-with-multitool-data")
    @app.route("/casestrainer/api/confirmed_with_multitool_data")
    def confirmed_with_multitool_data():
        """API endpoint for citations confirmed with multi-tool verification."""
        try:
            # Get citations from database
            conn = sqlite3.connect(DATABASE_FILE)
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM citations WHERE verified = 1")
            citations = cursor.fetchall()
            conn.close()

            # Format citations for response
            formatted_citations = []
            for citation in citations:
                formatted_citations.append(
                    {
                        "citation": citation[1],
                        "found": True,
                        "found_case_name": citation[2],
                        "confidence": citation[3],
                        "source": citation[4],
                        "url": citation[5],
                        "explanation": citation[6],
                    }
                )

            return jsonify(
                {"citations": formatted_citations, "total": len(formatted_citations)}
            )
        except Exception as e:
            logger.error(f"Error in confirmed_with_multitool_data: {str(e)}")
            logger.error(f"Error details: {traceback.format_exc()}")
            return jsonify({"error": str(e)}), 500

    @app.route("/api/processing_progress", methods=["GET"])
    def processing_progress():
        """Get the current processing progress."""
        return jsonify(processing_state)

    @app.route("/casestrainer/")
    @app.route("/casestrainer")
    def serve_vue_index():
        # Serve the Vue app's index.html directly
        vue_dist_dir = os.path.abspath(
            os.path.join(os.path.dirname(__file__), "..", "casestrainer-vue", "dist")
        )
        return send_from_directory(vue_dist_dir, "index.html")

    @app.route("/casestrainer/<path:path>")
    def serve_vue_assets(path):
        vue_dist_dir = os.path.abspath(
            os.path.join(os.path.dirname(__file__), "..", "casestrainer-vue", "dist")
        )
        response = send_from_directory(vue_dist_dir, path)
        response.headers["Cache-Control"] = (
            "no-store, no-cache, must-revalidate, max-age=0, post-check=0, pre-check=0"
        )
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "-1"
        response.headers["Last-Modified"] = (
            f"{datetime.datetime.now().strftime('%a, %d %b %Y %H:%M:%S GMT')}"
        )
        return response

    @app.route("/api/validate_citations", methods=["POST"])
    def validate_citations():
        """API endpoint for validating citations."""
        try:
            # Get the request data
            data = request.get_json()
            if not data or "text" not in data:
                return jsonify({"error": "No text provided"}), 400

            # Extract citations from the text (now returns list of dicts)
            extracted_citations = extract_citations_from_text(data["text"])

            # Verify each citation using the citation_text field
            results = []
            for citation in extracted_citations:
                verification = verify_citation(citation["citation_text"])
                # Combine extraction metadata and verification result
                results.append({"extraction": citation, "verification": verification})

            return jsonify({"citations": results, "total": len(results)})
        except Exception as e:
            logger.error(f"Error in validate_citations: {str(e)}")
            logger.error(f"Error details: {traceback.format_exc()}")
            return jsonify({"error": str(e)}), 500

    @app.after_request
    def after_request(response):
        """Clean up after each request."""
        # Clean up thread-local storage
        if hasattr(thread_local, "data"):
            del thread_local.data
        return response

    @app.errorhandler(404)
    def not_found(error):
        """Handle 404 errors."""
        return jsonify({"error": "Not found"}), 404

    @app.errorhandler(500)
    def server_error(error):
        """Handle 500 errors."""
        logger.error(f"Server error: {str(error)}")
        logger.error(f"Error details: {traceback.format_exc()}")
        return jsonify({"error": "Internal server error"}), 500

    return app


# Create the Flask app instance
app = create_app()


# Helper functions
def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


from citation_utils import (
    verify_citation,
    extract_citations_from_file,
    extract_all_citations,  # Use the new unified extraction function
    # ...
)


def extract_citations_from_file(filepath):
    """Extract citations from a file using the appropriate method based on file type."""
    start_time = time.time()
    try:
        file_extension = filepath.split(".")[-1].lower()
        if file_extension == "pdf":
            logger.info("Detected PDF file, using PyPDF2 for extraction")
            import PyPDF2

            try:
                pdf = PyPDF2.PdfReader(open(filepath, "rb"))
                text = ""
                for i in range(pdf.numPages):
                    page_text = pdf.getPage(i).extractText()
                    text += page_text
                    logger.info(f"Page {i+1} extracted: {len(page_text)} characters")
            except Exception as pdf_error:
                logger.error(f"Error reading PDF: {str(pdf_error)}")
                raise
        elif filepath.endswith((".doc", ".docx")):
            logger.info("Detected Word document, using python-docx for extraction")
            import docx

            try:
                doc = docx.Document(filepath)
                logger.info(f"Word document has {len(doc.paragraphs)} paragraphs")
                paragraphs = [paragraph.text for paragraph in doc.paragraphs]
                text = "\n".join(paragraphs)
                logger.info(
                    f"Extracted {len(paragraphs)} paragraphs with total {len(text)} characters"
                )
            except Exception as docx_error:
                logger.error(f"Error reading Word document: {str(docx_error)}")
                raise
        else:
            logger.info(f"Using standard text file reading for {file_extension} file")
            try:
                with open(filepath, "r", encoding="utf-8") as file:
                    text = file.read()
                    logger.info(f"Read {len(text)} characters from text file")
            except UnicodeDecodeError:
                logger.warning("UTF-8 decoding failed, trying with latin-1 encoding")
                with open(filepath, "r", encoding="latin-1") as file:
                    text = file.read()
                    logger.info(
                        f"Read {len(text)} characters from text file using latin-1 encoding"
                    )
        extraction_time = time.time() - start_time
        logger.info(
            f"File content extraction completed in {extraction_time:.2f} seconds"
        )
        logger.info(f"Extracted text sample (first 200 chars): {text[:200]}...")

        # Now extract citations from the text
        logger.info("Calling extract_citations_from_text with the extracted content")
        return extract_citations_from_text(text)
    except Exception as e:
        logger.error(f"Error extracting citations from file: {str(e)}")
        logger.error(traceback.format_exc())
        return []

    # REMOVED: extract_citations_from_text is now imported from citation_utils.py
    raise NotImplementedError(
        "extract_citations_from_text has moved to citation_utils.py. Please import from there."
    )
    """Extract citations from text using eyecite and return full metadata."""
    logger.info(
        f"[DEBUG] Starting extract_citations_from_text with text of length {len(text)}"
    )
    logger.info(f"[DEBUG] Text sample: {text[:300]}...")
    start_time = time.time()

    try:
        from eyecite import get_citations
        from eyecite.tokenizers import AhocorasickTokenizer

        logger.info("Successfully imported eyecite libraries")

        # Use AhocorasickTokenizer
        tokenizer = None
        try:
            logger.info("Attempting to initialize AhocorasickTokenizer")
            tokenizer_start = time.time()
            tokenizer = AhocorasickTokenizer()
            logger.info(
                f"AhocorasickTokenizer initialized in {time.time() - tokenizer_start:.2f} seconds"
            )
        except Exception as tokenizer_error:
            logger.warning(
                f"Failed to initialize AhocorasickTokenizer: {str(tokenizer_error)}"
            )
            logger.warning("Will fall back to default tokenizer")
            tokenizer = None

        logger.info("Starting citation extraction with eyecite")
        extraction_start = time.time()
        if tokenizer:
            logger.info("Using AhocorasickTokenizer for extraction")
            citations = get_citations(text, tokenizer=tokenizer)
        else:
            logger.info("Using default tokenizer for extraction")
            citations = get_citations(text)

        extraction_time = time.time() - extraction_start
        logger.info(f"Citation extraction completed in {extraction_time:.2f} seconds")
        logger.info(f"Found {len(citations)} citations in the text")

        logger.info("Processing citation objects to extract metadata")
        metadata_start = time.time()
        citation_dicts = []
        unique_citations = {}

        for i, citation in enumerate(citations):
            citation_text = citation.matched_text()
            logger.info(f"Processing citation {i+1}/{len(citations)}: {citation_text}")

            # Get case name if available
            case_name = None
            try:
                if (
                    hasattr(citation, "metadata")
                    and citation.metadata
                    and hasattr(citation.metadata, "case_name")
                ):
                    case_name = citation.metadata.case_name
                    logger.info(f"Found case name: {case_name}")
                else:
                    logger.info("No case name found in metadata")
            except Exception as name_error:
                logger.warning(f"Error extracting case name: {str(name_error)}")
                pass

            # Extract other metadata
            reporter = getattr(citation, "reporter", None)
            volume = getattr(citation, "volume", None)
            page = getattr(citation, "page", None)
            year = getattr(citation, "year", None)
            court = getattr(citation, "court", None)

            logger.info(
                f"Citation metadata - Reporter: {reporter}, Volume: {volume}, Page: {page}, Year: {year}, Court: {court}"
            )

            # Extract context (100 characters on either side of the citation)
            try:
                # Find the position of the citation in the text
                citation_pos = text.find(citation_text)
                if citation_pos != -1:
                    # Get 100 characters before and after the citation
                    start_pos = max(0, citation_pos - 100)
                    end_pos = min(len(text), citation_pos + len(citation_text) + 100)
                    context = text[start_pos:end_pos]
                    logger.info(
                        f"Extracted context of length {len(context)} for citation"
                    )
                else:
                    context = ""
                    logger.warning(
                        f"Could not find citation '{citation_text}' in text for context extraction"
                    )
            except Exception as context_error:
                logger.warning(f"Error extracting context: {str(context_error)}")
                context = ""

            # Create a unique key for deduplication (using volume, reporter, and page if available)
            citation_key = (
                f"{volume or ''}|{reporter or ''}|{page or ''}|{citation_text}"
            )

            # Create or update the citation dictionary
            if citation_key in unique_citations:
                # Add this context to the existing citation
                unique_citations[citation_key]["contexts"].append(
                    {"text": context, "citation_text": citation_text}
                )
                logger.info(f"Added context to existing citation: {citation_key}")
            else:
                # Create a new citation entry
                citation_dict = {
                    "text": citation_text,
                    "name": case_name or "Unknown Case",
                    "valid": None,  # Will be filled in by the API endpoint
                    "contexts": [{"text": context, "citation_text": citation_text}],
                    "metadata": {
                        "reporter": reporter,
                        "volume": volume,
                        "page": page,
                        "year": year,
                        "court": court,
                    },
                }
                unique_citations[citation_key] = citation_dict
                logger.info(f"Added new citation: {citation_key}")

        # Convert the dictionary of unique citations back to a list
        citation_dicts = list(unique_citations.values())
        logger.info(
            f"After deduplication, {len(citation_dicts)} unique citations remain"
        )

        metadata_time = time.time() - metadata_start
        logger.info(f"Metadata extraction completed in {metadata_time:.2f} seconds")

        total_time = time.time() - start_time
        logger.info(
            f"Total citation extraction process completed in {total_time:.2f} seconds"
        )
        logger.info(f"Returning {len(citation_dicts)} citation dictionaries")

        return citation_dicts
    except Exception as e:
        logger.error(f"Error extracting citations from text: {str(e)}")
        logger.error(traceback.format_exc())
        logger.error(f"Error occurred after processing text of length {len(text)}")
        # Log a sample of the text that caused the error
        if len(text) > 0:
            sample_size = min(500, len(text))
            logger.error(
                f"Text sample that caused the error (first {sample_size} chars): {text[:sample_size]}"
            )
        return []


def initialize_session_data():
    """Initialize session with sample citation data if it's empty."""
    from flask import session

    if "user_citations" not in session or not session["user_citations"]:
        logger.info("Initializing session with sample citation data")
        session["user_citations"] = [
            {
                "citation": "347 U.S. 483",
                "found": True,
                "found_case_name": "Brown v. Board of Education",
                "confidence": 0.95,
                "source": "Multi-source Verification",
                "url": "https://scholar.google.com/scholar_case?case=12120372216939101759",
                "explanation": "The landmark case Brown v. Board of Education (347 U.S. 483) established that separate educational facilities are inherently unequal.",
            },
            {
                "citation": "410 U.S. 113",
                "found": True,
                "found_case_name": "Roe v. Wade",
                "confidence": 0.92,
                "source": "CourtListener",
                "cl_id": "12345",
                "url": "https://www.courtlistener.com/opinion/12345",
                "explanation": "The Court's decision in Roe v. Wade (410 U.S. 113) recognized a woman's right to choose.",
            },
            {
                "citation": "5 U.S. 137",
                "found": True,
                "found_case_name": "Marbury v. Madison",
                "confidence": 0.88,
                "source": "Multi-source Verification",
                "url": "https://caselaw.findlaw.com/us-supreme-court/5/137.html",
                "explanation": "Marbury v. Madison (5 U.S. 137) established the principle of judicial review.",
            },
            {
                "citation": "123 U.S. 456",
                "found": False,
                "found_case_name": "Smith v. Jones",
                "confidence": 0.0,
                "source": "Not found",
                "explanation": "Citation not found in any source.",
            },
        ]


def get_ip_address():
    """Get the server's IP address for logging purposes"""
    try:
        # Get the server's IP address
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        s.connect(("8.8.8.8", 80))
        ip = s.getsockname()[0]
        s.close()
        return ip
    except Exception as e:
        logger.error(f"Error getting IP address: {e}")
        return "unknown"


if __name__ == "__main__":
    # Get command line arguments
    import argparse

    parser = argparse.ArgumentParser(
        description="Run CaseStrainer with Vue.js frontend"
    )
    # WARNING: Do not run app.run() here if using Waitress or another WSGI server for production.
    # Only use this for development/testing. Comment out or remove app.run() if using Waitress.
    parser.add_argument(
        "--host",
        default="0.0.0.0",
        help="Host to bind to (use 0.0.0.0 for Nginx proxy access)",
    )
    parser.add_argument(
        "--port",
        type=int,
        default=5000,
        help="Port to bind to (use 5000 for Nginx proxy access)",
    )
    parser.add_argument("--debug", action="store_true", help="Run in debug mode")
    parser.add_argument(
        "--use-waitress",
        action="store_true",
        help="Use Waitress WSGI server (production mode)",
    )
    parser.add_argument(
        "--env",
        choices=["development", "production"],
        default="production",
        help="Environment to run in",
    )
    parser.add_argument(
        "--threads", type=int, default=10, help="Number of threads for the server"
    )
    args = parser.parse_args()

    # Verify host and port settings
    if args.host != "0.0.0.0":
        logger.warning(
            f"Host set to {args.host} - this may prevent Nginx proxy access. Recommended: 0.0.0.0"
        )

    if args.port != 5000:
        logger.warning(
            f"Port set to {args.port} - this may prevent Nginx proxy access. Recommended: 5000"
        )

    # Set environment
    os.environ["FLASK_ENV"] = args.env

    # Check if we should run with Waitress (production) or Flask's dev server
    use_waitress = args.use_waitress or os.environ.get(
        "USE_WAITRESS", "True"
    ).lower() in ("true", "1", "t")

    # Check if port is already in use
    import socket

    sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    result = sock.connect_ex(("127.0.0.1", args.port))
    if result == 0:
        logger.warning(
            f"Port {args.port} is already in use. The application may not start correctly."
        )
        logger.warning(
            "Use start_for_nginx.bat or start_casestrainer.bat to automatically handle port conflicts."
        )
    sock.close()

    # Log startup information
    logger.info(f"Starting CaseStrainer with Vue.js frontend")
    logger.info(f"Host: {args.host}, Port: {args.port}, Environment: {args.env}")
    logger.info(f"Using Waitress: {use_waitress}")
    logger.info(f"External access URL: https://wolf.law.uw.edu/casestrainer/")
    logger.info(f"Local access URL: http://127.0.0.1:{args.port}/")

    if use_waitress:
        try:
            from waitress import serve

            logger.info("Starting with Waitress WSGI server (production mode)")

            # Now 'app' is defined and can be used
            serve(app, host=args.host, port=args.port, threads=args.threads)
        except ImportError:
            logger.warning("Waitress not installed. Installing now...")
            try:
                import subprocess

                subprocess.check_call(
                    [sys.executable, "-m", "pip", "install", "waitress"]
                )
                logger.info("Waitress installed. Please restart the application.")
                sys.exit(1)
            except Exception as e:
                logger.error(f"Failed to install Waitress: {e}")
                logger.info("Falling back to Flask development server")
                app.run(debug=args.debug, host=args.host, port=args.port)
    else:
        logger.info("Starting with Flask development server")
        app.run(debug=args.debug, host=args.host, port=args.port)
